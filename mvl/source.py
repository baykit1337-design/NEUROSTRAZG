"""Единая точка чтения исходников: epub, docx, txt, md.

Раньше каждая вкладка читала файлы по-своему, и `.epub` в двух местах
открывался как текст — а это ZIP-архив, отсюда и мусор вида `??v????A`.
Теперь формат распознаётся в одном месте.

Разбор epub берётся из `split_book.epub_chapters` — он уже обкатан и здесь
не переписывается.
"""

from __future__ import annotations

import logging
import re
from dataclasses import dataclass
from pathlib import Path

from .booksplit import read_chapters as read_book_chapters
from .word import split_paragraphs

log = logging.getLogger(__name__)

EPUB = (".epub",)
DOCX = (".docx",)
PLAIN = (".txt", ".md")
#: Всё, что вкладки умеют читать.
READABLE = EPUB + DOCX + PLAIN
#: Для системного диалога выбора файлов.
DIALOG_FILTER = "*.epub *.docx *.txt *.md"


class SourceError(Exception):
    """Файл не удалось прочитать."""


@dataclass
class Chapter:
    """Одна глава: заголовок и абзацы. Источник может дать сразу много."""

    title: str
    paragraphs: list[str]
    source: str = ""

    @property
    def text(self) -> str:
        return "\n\n".join(self.paragraphs)

    @property
    def size(self) -> int:
        return sum(len(p) for p in self.paragraphs)


def is_readable(path: Path) -> bool:
    return path.suffix.lower() in READABLE


def is_multi_chapter(path: Path) -> bool:
    """Даёт ли файл несколько глав сразу. Пока это только epub."""
    return path.suffix.lower() in EPUB


def read_paragraphs(path: Path) -> list[str]:
    """Абзацы одного плоского файла: .txt, .md или .docx.

    Для epub не годится — там глав много, нужен `load_chapters`.
    """
    suffix = path.suffix.lower()
    if suffix in DOCX:
        try:
            from docx import Document
        except ImportError as exc:
            raise SourceError("Для .docx нужен python-docx") from exc
        try:
            document = Document(str(path))
        except Exception as exc:
            raise SourceError(f"{type(exc).__name__}: {exc}") from exc
        return [p.text.strip() for p in document.paragraphs if p.text.strip()]

    if suffix in EPUB:
        # Читать архив как текст нельзя — вернём главы склеенными.
        return [p for chapter in load_chapters(path) for p in chapter.paragraphs]

    try:
        return split_paragraphs(path.read_text(encoding="utf-8", errors="replace"))
    except OSError as exc:
        raise SourceError(f"{type(exc).__name__}: {exc}") from exc


def load_chapters(path: str | Path) -> list[Chapter]:
    """Главы из файла. epub даёт много, остальные форматы — одну."""
    file_path = Path(path).expanduser()
    if not file_path.is_file():
        raise SourceError(f"Файл не найден: {file_path}")

    suffix = file_path.suffix.lower()
    if suffix not in READABLE:
        raise SourceError(
            f"Не умею читать {suffix or 'файл без расширения'} — "
            "нужен .epub, .docx, .txt или .md"
        )

    if suffix in EPUB:
        # epub — ZIP-архив: распаковываем и разбираем по OPF.
        try:
            chapters = read_book_chapters(file_path)
        except Exception as exc:
            raise SourceError(f"{file_path.name}: {exc}") from exc
        return [
            Chapter(title=c.title, paragraphs=split_paragraphs(c.text), source=str(file_path))
            for c in chapters
        ]

    return [
        Chapter(
            title=clean_title(file_path.stem),
            paragraphs=read_paragraphs(file_path),
            source=str(file_path),
        )
    ]


#: Порядковый номер файла «0010 - » в начале имени. В заголовок главы он не
#: идёт: это служебная нумерация WebToEpub, нужная только для сортировки.
SEQ_PREFIX = re.compile(r"^\s*\d+\s*[-–—]\s*")


def clean_title(stem: str) -> str:
    """Заголовок главы из имени файла, без порядкового номера."""
    return SEQ_PREFIX.sub("", stem.strip()).strip() or stem.strip()


def collect_sources(targets) -> list[Path]:
    """Разворачивает выбранное в список файлов.

    Принимает пути к файлам и к папкам вперемешку: папка раскрывается в
    свои читаемые файлы, файл берётся как есть.
    """
    if isinstance(targets, (str, Path)):
        targets = [targets]

    files: list[Path] = []
    seen: set[str] = set()

    for target in targets:
        path = Path(str(target)).expanduser()
        if path.is_dir():
            found = [p for p in sorted(path.iterdir()) if p.is_file() and is_readable(p)]
        elif path.is_file():
            # Расширение проверяем и у одиночного файла: без этого epub
            # уходил в текстовое чтение и превращался в мусор.
            if not is_readable(path):
                raise SourceError(
                    f"{path.name}: нужен .epub, .docx, .txt или .md"
                )
            found = [path]
        else:
            raise SourceError(f"Не найдено: {path}")

        for item in found:
            key = str(item.resolve())
            if key not in seen:
                seen.add(key)
                files.append(item)

    if not files:
        raise SourceError("Не нашлось файлов .epub, .docx, .txt или .md")
    return files
