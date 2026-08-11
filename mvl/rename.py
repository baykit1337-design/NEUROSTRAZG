"""Переименование и деление глав.

Имена от WebToEpub приходят вперемешку: порядковый номер файла не совпадает
с номером главы, а первой идёт служебная «Информация». Здесь имя
раскладывается на части, собирается заново по галочкам, а сами главы при
необходимости режутся на несколько файлов.

Оригиналы не трогаются никогда — результат пишется в новую папку.
"""

from __future__ import annotations

import logging
import re
import threading
from dataclasses import dataclass, field
from pathlib import Path

from .booksplit import Cancelled, safe_name
from .word import SCENE_BREAK, Style, split_paragraphs

log = logging.getLogger(__name__)

READABLE = (".txt", ".md", ".docx")

#: Слова, которыми помечают главу. Нужны, чтобы отличить номер главы от
#: порядкового номера файла.
CHAPTER_WORD = r"(?:Глава|ГЛАВА|глава|Chapter|CHAPTER|chapter|Часть|Part)"

#: Порядковый номер файла в начале имени: «0010 - ».
SEQ_RE = re.compile(r"^\s*(\d+)\s*[-–—]\s*(.*)$", re.S)

#: Номер главы с необязательной приставкой и необязательным номером части.
NUMBER_RE = re.compile(
    rf"^\s*(?:{CHAPTER_WORD}\s*)?"
    r"(?P<number>\d+)"
    r"(?:[.,](?P<part>\d+))?"
    r"(?:\s*(?:[.:\-–—]|\s)\s*(?P<title>.*?))?\s*$",
    re.S,
)

#: Совместимость: та же регулярка одной строкой — её показывает интерфейс
#: как отправную точку для своего варианта.
DEFAULT_PATTERN = NUMBER_RE.pattern

DEFAULT_PREFIX = "Глава"
DEFAULT_SEPARATOR = ": "
#: Разделители перед названием, предлагаемые в интерфейсе.
SEPARATORS = (": ", ". ", " — ", " - ")

#: Windows не разрешает эти символы в именах файлов. Просто выбрасывать их
#: нельзя: из «Глава 201: Конец» вышло бы «Глава 201 Конец», и разделитель
#: пропал бы совсем. Поэтому подставляем читаемый эквивалент.
FORBIDDEN_MAP = {
    ":": " -", "/": "-", "\\": "-", "|": "-",
    "*": "", "?": "", '"': "'", "<": "(", ">": ")",
}


def safe_filename(name: str) -> str:
    """Имя файла, пригодное для Windows, с сохранением читаемости."""
    for bad, good in FORBIDDEN_MAP.items():
        name = name.replace(bad, good)
    name = re.sub(r"\s+", " ", name).strip(" .")
    return safe_name(name)


def has_forbidden(text: str) -> bool:
    """Есть ли в строке символы, запрещённые в именах файлов."""
    return any(bad in text for bad in FORBIDDEN_MAP)


class RenameError(Exception):
    """Папку не удалось разобрать."""


@dataclass
class NameParts:
    seq: int | None = None
    number: int | None = None
    part: int | None = None
    title: str = ""

    @property
    def service(self) -> bool:
        """Номера главы нет — файл служебный («Информация», «Обложка»)."""
        return self.number is None


def parse_name(stem: str, pattern: str | None = None) -> NameParts:
    """Раскладывает имя файла (без расширения) на составляющие.

    Разбор идёт в два шага, и это принципиально: сначала отрезается
    порядковый номер файла, и только в остатке ищется номер главы. Иначе
    `0001 - Информация` превращается в «главу 1», хотя это служебный файл
    без номера вовсе.

    Порядковый номер разбирается, но в новое имя не попадает — он нужен
    только чтобы понять исходный порядок.
    """
    if pattern:
        try:
            match = re.match(pattern, stem)
        except re.error as exc:
            raise RenameError(f"Неверное регулярное выражение: {exc}") from exc
        if not match:
            return NameParts(seq=_leading_int(stem), title=stem.strip())
        groups = match.groupdict()
        return NameParts(
            seq=_int(groups.get("seq")),
            number=_int(groups.get("number")),
            part=_int(groups.get("part")),
            title=(groups.get("title") or "").strip(),
        )

    seq = None
    rest = stem.strip()
    head = SEQ_RE.match(rest)
    if head:
        seq, rest = int(head.group(1)), head.group(2).strip()

    match = NUMBER_RE.match(rest)
    if not match:
        # Номера главы нет — служебный файл («Информация», «Обложка»).
        return NameParts(seq=seq if seq is not None else _leading_int(stem),
                         title=rest or stem.strip())

    return NameParts(
        seq=seq,
        number=_int(match.group("number")),
        part=_int(match.group("part")),
        title=(match.group("title") or "").strip(),
    )


def _int(value) -> int | None:
    try:
        return int(value)
    except (TypeError, ValueError):
        return None


def _leading_int(stem: str) -> int | None:
    match = re.match(r"\s*(\d+)", stem)
    return int(match.group(1)) if match else None


@dataclass
class Chapter:
    """Одна глава: файл на диске плюс то, что вышло из его имени."""

    path: Path
    parts_of_name: NameParts
    size: int = 0
    part: int | None = None
    text_parts: list[str] = field(default_factory=list)

    @property
    def number(self) -> int | None:
        return self.parts_of_name.number

    @property
    def title(self) -> str:
        return self.parts_of_name.title

    @property
    def service(self) -> bool:
        return self.parts_of_name.service

    def as_dict(self) -> dict:
        return {
            "path": str(self.path),
            "name": self.path.name,
            "number": self.number,
            "part": self.part,
            "title": self.title,
            "size": self.size,
            "service": self.service,
        }


def read_paragraphs(path: Path) -> list[str]:
    """Абзацы файла. Понимает .txt, .md и .docx."""
    suffix = path.suffix.lower()
    if suffix == ".docx":
        try:
            from docx import Document
        except ImportError as exc:
            raise RenameError("Для .docx нужен python-docx") from exc
        document = Document(str(path))
        return [p.text.strip() for p in document.paragraphs if p.text.strip()]
    # UTF-8 с errors='replace' — битый файл не должен ронять обработку.
    return split_paragraphs(path.read_text(encoding="utf-8", errors="replace"))


def scan(folder: str | Path, pattern: str | None = None) -> list[Chapter]:
    """Читает папку и раскладывает имена. Сортировка — по номеру главы."""
    directory = Path(folder).expanduser()
    if not directory.is_dir():
        raise RenameError(f"Папка не найдена: {directory}")

    chapters: list[Chapter] = []
    for path in sorted(directory.iterdir(), key=lambda p: p.name.lower()):
        if not path.is_file() or path.suffix.lower() not in READABLE:
            continue
        parts = parse_name(path.stem, pattern)
        try:
            paragraphs = read_paragraphs(path)
        except Exception as exc:
            log.warning("Не прочитан %s: %s", path.name, exc)
            paragraphs = []
        chapters.append(
            Chapter(
                path=path,
                parts_of_name=parts,
                part=parts.part,
                size=sum(len(p) for p in paragraphs),
                text_parts=paragraphs,
            )
        )

    if not chapters:
        raise RenameError("В папке нет файлов .txt, .md или .docx")

    return sort_chapters(chapters)


def sort_chapters(chapters: list[Chapter]) -> list[Chapter]:
    """Сначала главы по номеру, служебные — в конец, но не теряются."""
    return sorted(
        chapters,
        key=lambda c: (
            c.service,
            c.number if c.number is not None else 0,
            c.part if c.part is not None else 0,
            c.parts_of_name.seq if c.parts_of_name.seq is not None else 0,
        ),
    )


# --------------------------------------------------------------- сборка имени


@dataclass
class NameFormat:
    """Три галочки из ТЗ плюс приставка и разделитель."""

    number: bool = True
    part: bool = True
    title: bool = True
    prefix: str = DEFAULT_PREFIX
    separator: str = DEFAULT_SEPARATOR

    @classmethod
    def from_dict(cls, data: dict | None) -> NameFormat:
        data = data or {}
        prefix = data.get("prefix")
        separator = data.get("separator")
        return cls(
            number=bool(data.get("number", True)),
            part=bool(data.get("part", True)),
            title=bool(data.get("title", True)),
            # Приставку можно очистить, поэтому проверяем на None, а не на «пусто».
            prefix=DEFAULT_PREFIX if prefix is None else str(prefix),
            separator=DEFAULT_SEPARATOR if separator is None else str(separator),
        )


def build_name(number: int | None, part: int | None, title: str, fmt: NameFormat) -> str:
    """Собирает новое имя главы по галочкам.

    Номер части подставляется только если он есть: у целой главы части нет,
    и включённая галочка ничего не добавляет.
    """
    head = ""
    if fmt.number and number is not None:
        head = f"{fmt.prefix} {number}".strip() if fmt.prefix else str(number)
        if fmt.part and part is not None:
            head = f"{head}.{part}"

    name = head
    if fmt.title and title:
        name = f"{head}{fmt.separator}{title}" if head else title

    # Все галочки сняты — пустое имя недопустимо, оставляем номер.
    if not name.strip():
        name = str(number) if number is not None else title
    return safe_filename(name.strip())


# ------------------------------------------------------------ деление на части


def split_into_parts(paragraphs: list[str], count: int) -> list[list[str]]:
    """Делит абзацы на `count` частей, максимально равных по числу символов.

    Режем только по границам абзацев. Разделитель сцен (`*`) не должен
    оказаться первым или последним абзацем части — такие границы сдвигаем.
    """
    if count < 2:
        return [list(paragraphs)]
    blocks = [p for p in paragraphs if p.strip()]
    if len(blocks) < count:
        # Абзацев меньше, чем частей — делить нечего.
        return [list(blocks)]

    lengths = [len(p) for p in blocks]
    total = sum(lengths)
    # Накопленная длина после каждого абзаца.
    cumulative: list[int] = []
    running = 0
    for length in lengths:
        running += length
        cumulative.append(running)

    cuts: list[int] = []
    for index in range(1, count):
        ideal = total * index / count
        # Граница — индекс абзаца, после которого режем.
        best = min(range(len(blocks)), key=lambda i: abs(cumulative[i] - ideal))
        cut = best + 1
        # Каждая часть должна быть непустой, границы строго возрастают.
        low = (cuts[-1] + 1) if cuts else 1
        high = len(blocks) - (count - index)
        cuts.append(max(low, min(cut, high)))

    cuts = _avoid_scene_breaks(blocks, cuts)

    parts: list[list[str]] = []
    start = 0
    for cut in [*cuts, len(blocks)]:
        parts.append(blocks[start:cut])
        start = cut
    return [p for p in parts if p]


def _avoid_scene_breaks(blocks: list[str], cuts: list[int]) -> list[int]:
    """Сдвигает границы так, чтобы `*` не открывал и не закрывал часть."""
    adjusted = []
    for index, cut in enumerate(cuts):
        low = (adjusted[-1] + 1) if adjusted else 1
        high = len(blocks) - (len(cuts) - index)
        for candidate in _nearby(cut, low, high):
            ends_with_break = SCENE_BREAK.match(blocks[candidate - 1])
            starts_with_break = SCENE_BREAK.match(blocks[candidate])
            if not ends_with_break and not starts_with_break:
                cut = candidate
                break
        adjusted.append(max(low, min(cut, high)))
    return adjusted


def _nearby(value: int, low: int, high: int):
    """Кандидаты в порядке удаления от исходной границы."""
    if low > high:
        return
    seen = set()
    for shift in range(0, high - low + 2):
        for candidate in (value - shift, value + shift):
            if low <= candidate <= high and candidate not in seen:
                seen.add(candidate)
                yield candidate


# ------------------------------------------------------------------- план работ


@dataclass
class PlanRow:
    """Строка предпросмотра: что было и что станет."""

    source: str
    old_name: str
    new_name: str
    number: int | None
    part: int | None
    title: str
    size: int
    service: bool
    paragraphs: list[str] = field(default_factory=list)

    def as_dict(self) -> dict:
        return {
            "source": self.source,
            "old_name": self.old_name,
            "new_name": self.new_name,
            "number": self.number,
            "part": self.part,
            "title": self.title,
            "size": self.size,
            "service": self.service,
        }


@dataclass
class RenameReport:
    output_dir: str = ""
    written: int = 0
    failed: int = 0
    failed_files: list[str] = field(default_factory=list)
    total: int = 0

    def as_dict(self) -> dict:
        return {
            "output_dir": self.output_dir,
            "written": self.written,
            "failed": self.failed,
            "failed_files": self.failed_files,
            "total": self.total,
        }


def make_plan(
    chapters: list[Chapter],
    fmt: NameFormat,
    splits: dict[str, int] | None = None,
    renumber_from: int | None = None,
    skip_service: bool = True,
) -> list[PlanRow]:
    """Строит предпросмотр «старое имя → новое имя».

    `splits` — сколько частей сделать из главы, ключ это путь к файлу.
    `renumber_from` — нумеровать заново подряд, игнорируя номер из имени.
    """
    splits = splits or {}
    rows: list[PlanRow] = []
    next_number = renumber_from

    for chapter in sort_chapters(chapters):
        if chapter.service and skip_service:
            continue

        if chapter.service:
            # Служебный файл переименованию не поддаётся — оставляем как есть.
            rows.append(
                PlanRow(
                    source=str(chapter.path),
                    old_name=chapter.path.name,
                    new_name=safe_filename(chapter.path.stem),
                    number=None, part=None, title=chapter.title,
                    size=chapter.size, service=True,
                    paragraphs=chapter.text_parts,
                )
            )
            continue

        if renumber_from is not None:
            number = next_number
            next_number += 1
        else:
            number = chapter.number

        count = max(1, int(splits.get(str(chapter.path), 1) or 1))
        if count > 1:
            pieces = split_into_parts(chapter.text_parts, count)
            for index, piece in enumerate(pieces, 1):
                rows.append(
                    PlanRow(
                        source=str(chapter.path),
                        old_name=chapter.path.name,
                        new_name=build_name(number, index, chapter.title, fmt),
                        number=number, part=index, title=chapter.title,
                        size=sum(len(p) for p in piece), service=False,
                        paragraphs=piece,
                    )
                )
        else:
            rows.append(
                PlanRow(
                    source=str(chapter.path),
                    old_name=chapter.path.name,
                    new_name=build_name(number, chapter.part, chapter.title, fmt),
                    number=number, part=chapter.part, title=chapter.title,
                    size=chapter.size, service=False,
                    paragraphs=chapter.text_parts,
                )
            )

    return rows


def apply_plan(
    rows: list[PlanRow],
    output_dir: Path,
    fmt: str = "txt",
    style: Style | None = None,
    on_progress=None,
    cancel: threading.Event | None = None,
) -> RenameReport:
    """Пишет результат в новую папку. Оригиналы не трогаются.

    Сбой на одном файле не останавливает остальные.
    """
    output_dir.mkdir(parents=True, exist_ok=True)
    report = RenameReport(output_dir=str(output_dir), total=len(rows))
    used: set[str] = set()

    for index, row in enumerate(rows, 1):
        if cancel is not None and cancel.is_set():
            raise Cancelled()

        suffix = fmt if fmt in ("txt", "docx") else "txt"
        name = _unique(safe_filename(row.new_name) or f"{index:04d}", suffix, used)
        try:
            target = output_dir / name
            if suffix == "docx":
                from .word import write_chapter

                write_chapter(target, "", "\n\n".join(row.paragraphs), style)
            else:
                target.write_text("\n\n".join(row.paragraphs) + "\n", encoding="utf-8")
            report.written += 1
        except Exception as exc:
            log.warning("Не записан %s: %s", name, exc)
            report.failed += 1
            report.failed_files.append(f"{name}: {type(exc).__name__}: {exc}")

        if on_progress:
            on_progress(index, len(rows))

    return report


def _unique(stem: str, suffix: str, used: set[str]) -> str:
    """Разводит совпадающие имена, чтобы файлы не затирали друг друга."""
    name = f"{stem}.{suffix}"
    if name.lower() not in used:
        used.add(name.lower())
        return name
    counter = 2
    while f"{stem} ({counter}).{suffix}".lower() in used:
        counter += 1
    name = f"{stem} ({counter}).{suffix}"
    used.add(name.lower())
    return name
