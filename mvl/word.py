"""Сборка .docx: общие настройки документа и вставка глав.

Используется вкладкой «Разбить книгу» (глава = отдельный файл) и будет
переиспользована конвертером в Word (все главы в одном документе).
"""

from __future__ import annotations

import re
from dataclasses import dataclass
from pathlib import Path

# Абзац из одних звёздочек — разделитель сцен, его сохраняем как есть.
SCENE_BREAK = re.compile(r"^[*\s]*\*[*\s]*$")


class DocxUnavailable(RuntimeError):
    """Не установлен python-docx."""


@dataclass
class Style:
    """Оформление документа. Значения по умолчанию — из ТЗ."""

    font: str = "Times New Roman"
    size: int = 12
    line_spacing: float = 1.5
    #: По ТЗ v4 красной строки по умолчанию нет — было 1.25 см.
    first_line_indent_cm: float = 0.0
    page_break_between_chapters: bool = True

    @classmethod
    def from_dict(cls, data: dict | None) -> Style:
        data = data or {}
        return cls(
            font=str(data.get("font") or cls.font),
            size=_number(data.get("size"), cls.size),
            line_spacing=_number(data.get("line_spacing"), cls.line_spacing),
            first_line_indent_cm=_number(
                data.get("first_line_indent_cm"), cls.first_line_indent_cm
            ),
            page_break_between_chapters=bool(
                data.get("page_break_between_chapters", cls.page_break_between_chapters)
            ),
        )


def _number(value, fallback):
    try:
        number = float(value)
    except (TypeError, ValueError):
        return fallback
    if number <= 0:
        return fallback
    return int(number) if isinstance(fallback, int) else number


def _docx():
    try:
        import docx
    except ImportError as exc:  # pragma: no cover — зависит от окружения
        raise DocxUnavailable(
            "Для вывода в .docx нужен python-docx: pip install python-docx"
        ) from exc
    return docx


def new_document(style: Style | None = None):
    """Пустой документ с применённым оформлением."""
    docx = _docx()
    from docx.shared import Cm, Pt

    style = style or Style()
    document = docx.Document()

    normal = document.styles["Normal"]
    normal.font.name = style.font
    normal.font.size = Pt(style.size)
    paragraph_format = normal.paragraph_format
    paragraph_format.line_spacing = style.line_spacing
    paragraph_format.first_line_indent = Cm(style.first_line_indent_cm)

    # Кириллица берёт шрифт из отдельного атрибута темы, без него Word
    # подставляет свой.
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    for attribute in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rfonts.set(_qn(attribute), style.font)

    return document


def _qn(tag: str) -> str:
    from docx.oxml.ns import qn

    return qn(tag)


# Разметка markdown: её надо отрисовывать, а не тащить символами в текст.
HEADING = re.compile(r"^(#{1,6})\s+(.*)$", re.S)
QUOTE = re.compile(r"^>\s?(.*)$", re.S)
RULE = re.compile(r"^\s*(?:-{3,}|_{3,})\s*$")
#: Жирный, курсив, моноширинный и ссылка вида [текст](адрес).
INLINE = re.compile(
    r"(\*\*.+?\*\*|__.+?__|\*[^*\n]+?\*|_[^_\n]+?_|`[^`\n]+?`|\[[^\]]+?\]\([^)]*?\))",
    re.S,
)
LINK = re.compile(r"^\[([^\]]+?)\]\(([^)]*?)\)$", re.S)


def add_runs(paragraph, text: str) -> None:
    """Разбирает строчную разметку в отдельные run-ы с оформлением."""
    for token in INLINE.split(text):
        if not token:
            continue
        bold = italic = mono = False
        body = token

        if token.startswith("**") and token.endswith("**") and len(token) > 4:
            bold, body = True, token[2:-2]
        elif token.startswith("__") and token.endswith("__") and len(token) > 4:
            bold, body = True, token[2:-2]
        elif token.startswith("*") and token.endswith("*") and len(token) > 2:
            italic, body = True, token[1:-1]
        elif token.startswith("_") and token.endswith("_") and len(token) > 2:
            italic, body = True, token[1:-1]
        elif token.startswith("`") and token.endswith("`") and len(token) > 2:
            mono, body = True, token[1:-1]
        else:
            link = LINK.match(token)
            if link:
                # Адрес в бумажном тексте не нужен — оставляем подпись.
                body = link.group(1)

        run = paragraph.add_run(body)
        run.bold = bold
        run.italic = italic
        if mono:
            run.font.name = "Courier New"


def _alignment(name: str):
    """Название выравнивания → константа Word."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    return {
        "left": WD_ALIGN_PARAGRAPH.LEFT,
        "center": WD_ALIGN_PARAGRAPH.CENTER,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }.get(name, WD_ALIGN_PARAGRAPH.LEFT)


def add_blocks(document, blocks, style: Style | None = None, prep=None) -> int:
    """Вставляет подготовленные блоки в документ.

    Выравнивание и отступ проставляются каждому абзацу явно: исходное
    форматирование из epub или html наследовать нельзя, иначе часть текста
    приезжает по центру вперемешку с обычными абзацами.
    """
    from .textprep import KIND_SCENE, KIND_SYSTEM, PrepOptions

    style = style or Style()
    prep = prep or PrepOptions()
    body_align = _alignment(prep.align)
    indent = _cm(prep.first_line_indent_cm)

    added = 0
    for block in blocks:
        if block.kind == KIND_SCENE and not block.text:
            # «Пустая строка» — разделитель убран, остаётся отбивка.
            paragraph = document.add_paragraph()
            paragraph.alignment = body_align
            paragraph.paragraph_format.first_line_indent = 0
            added += 1
            continue

        if block.kind == KIND_SCENE:
            paragraph = document.add_paragraph()
            paragraph.alignment = _alignment("center")
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.add_run(block.text)
            added += 1
            continue

        # Блочная разметка markdown: заголовок и цитата — не просто текст.
        heading = HEADING.match(block.text)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            added += 1
            continue

        quote = QUOTE.match(block.text)
        if quote:
            paragraph = document.add_paragraph()
            paragraph.alignment = body_align
            paragraph.paragraph_format.left_indent = _cm(1.0)
            paragraph.paragraph_format.first_line_indent = 0
            add_runs(paragraph, quote.group(1).strip())
            for run in paragraph.runs:
                run.italic = True
            added += 1
            continue

        # Системные сообщения выравниваются как обычный текст.
        paragraph = document.add_paragraph()
        paragraph.alignment = body_align
        paragraph.paragraph_format.first_line_indent = indent
        add_runs(paragraph, block.text)
        if block.kind == KIND_SYSTEM and prep.italic_system:
            for run in paragraph.runs:
                run.italic = True
        added += 1
    return added


def add_paragraphs(document, text: str, style: Style | None = None, prep=None) -> int:
    """Добавляет текст главы абзацами, разбирая markdown.

    Возвращает число добавленных абзацев.
    """
    from .textprep import KIND_TEXT, Block, PrepOptions

    style = style or Style()
    prep = prep or PrepOptions()
    body_align = _alignment(prep.align)
    indent = _cm(prep.first_line_indent_cm)

    added = 0
    for chunk in split_paragraphs(text):
        # Разделитель сцен проверяем первым: «***» — это сцена, а не разметка.
        if SCENE_BREAK.match(chunk) or RULE.match(chunk):
            added += add_blocks(document, [Block(chunk.strip(), "scene")], style, prep)
            continue

        heading = HEADING.match(chunk)
        if heading:
            level = min(len(heading.group(1)), 4)
            document.add_heading(heading.group(2).strip(), level=level)
            added += 1
            continue

        quote = QUOTE.match(chunk)
        if quote:
            paragraph = document.add_paragraph()
            paragraph.alignment = body_align
            paragraph.paragraph_format.left_indent = _cm(1.0)
            paragraph.paragraph_format.first_line_indent = 0
            add_runs(paragraph, quote.group(1).strip())
            for run in paragraph.runs:
                run.italic = True
            added += 1
            continue

        paragraph = document.add_paragraph()
        paragraph.alignment = body_align
        paragraph.paragraph_format.first_line_indent = indent
        add_runs(paragraph, chunk)
        added += 1

    _ = KIND_TEXT
    return added


def _cm(value: float):
    from docx.shared import Cm

    return Cm(value)


def split_paragraphs(text: str) -> list[str]:
    """Режет текст на абзацы: пустая строка — граница."""
    if not text:
        return []
    blocks = re.split(r"\n\s*\n", text.replace("\r\n", "\n").replace("\r", "\n"))
    return [b.strip() for b in blocks if b.strip()]


def add_chapter(document, title: str, text: str, style: Style | None = None, prep=None) -> None:
    """Заголовок главы стилем Heading 1, дальше подготовленные абзацы.

    Текст всегда проходит через общую подготовку: дубль названия убирается,
    разделители схлопываются, пустые абзацы выбрасываются.
    """
    from .textprep import PrepOptions, prepare

    prep = prep or PrepOptions()
    if title:
        document.add_heading(title, level=1)
    add_blocks(document, prepare(split_paragraphs(text), title, prep), style, prep)


def write_chapter(
    path: Path, title: str, text: str, style: Style | None = None, prep=None
) -> None:
    """Одна глава — один .docx."""
    style = style or Style()
    document = new_document(style)
    add_chapter(document, title, text, style, prep)
    document.save(str(path))
