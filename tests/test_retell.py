"""Пересказ и выгрузка на базе реестра (3.5 ТЗ NEUROSTRAZH)."""

from __future__ import annotations

import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core.registry import Entity, Event, Link, Registry  # noqa: E402
from ops import docs, retell  # noqa: E402


def book() -> Registry:
    registry = Registry()
    registry.add_entity(Entity(name="Ли Вэй", type="персонаж", first_chapter=1))
    registry.add_entity(Entity(name="Чжан", type="персонаж", first_chapter=1))
    registry.add_entity(Entity(name="Меч Зари", type="предмет", first_chapter=2))
    registry.add_link(Link(source="li-vej", target="mech-zari", type="владеет"))
    registry.add_event(Event(chapter=1, type="встреча", actor="Ли Вэй",
                             object="Чжан", quote="Они встретились у ворот."))
    registry.add_event(Event(chapter=2, type="получение", actor="Ли Вэй",
                             object="Меч Зари", quote="Меч лёг в руку."))
    registry.add_event(Event(chapter=2, type="использование", actor="Ли Вэй",
                             object="Меч Зари"))
    return registry


class TestChapterRetell(unittest.TestCase):
    """Пересказ по главам — из фактов, без новых запросов."""

    def setUp(self):
        self.registry = book()

    def test_one_entry_per_chapter(self):
        items = retell.chapters(self.registry)
        self.assertEqual([i.chapter for i in items], [1, 2])

    def test_events_become_readable_lines(self):
        items = retell.chapters(self.registry)
        self.assertIn("Ли Вэй — встреча — Чжан", items[0].lines)

    def test_identifiers_are_replaced_with_names(self):
        """В событиях лежат идентификаторы, а читать надо имена."""
        for item in retell.chapters(self.registry):
            for line in item.lines:
                self.assertNotIn("li-vej", line)

    def test_new_names_are_marked(self):
        items = retell.chapters(self.registry)
        self.assertIn("Ли Вэй", items[0].new_names)
        self.assertIn("Меч Зари", items[1].new_names)

    def test_repeated_events_are_not_doubled(self):
        registry = Registry()
        for _ in range(3):
            registry.add_event(Event(chapter=5, type="встреча", actor="А",
                                     object="Б"))
        self.assertEqual(len(retell.chapters(registry)[0].lines), 1)

    def test_chapter_is_not_a_protocol(self):
        """Двадцать строк на главу — это уже не пересказ."""
        registry = Registry()
        for n in range(20):
            registry.add_event(Event(chapter=1, type=f"событие{n}", actor="А"))
        self.assertLessEqual(len(retell.chapters(registry)[0].lines),
                             retell.EVENTS_PER_CHAPTER)

    def test_events_without_a_chapter_are_skipped(self):
        registry = Registry()
        registry.add_event(Event(chapter=None, type="встреча", actor="А"))
        self.assertEqual(retell.chapters(registry), [])

    def test_text_has_chapter_headings(self):
        text = retell.chapters_text(self.registry)
        self.assertIn("# Глава 1", text)
        self.assertIn("Впервые появляются:", text)

    def test_empty_registry_gives_empty_text(self):
        self.assertEqual(retell.chapters_text(Registry()), "")


class TestAnnotation(unittest.TestCase):
    """Аннотация книги — один запрос."""

    class FakeClient:
        def __init__(self, answer="Аннотация книги."):
            self.answer = answer
            self.calls = 0
            self.prompt = ""

        def generate(self, prompt, json_only=True, model=""):
            self.calls += 1
            self.prompt = prompt
            return self.answer

    def test_one_request_only(self):
        client = self.FakeClient("текст" * 250)
        retell.annotation(book(), client)
        self.assertEqual(client.calls, 1)

    def test_facts_reach_the_prompt(self):
        client = self.FakeClient("текст" * 250)
        retell.annotation(book(), client)
        self.assertIn("Ли Вэй", client.prompt)
        self.assertIn("глава 1", client.prompt)

    def test_length_is_reported_against_the_spec(self):
        client = self.FakeClient("а" * 1200)
        result = retell.annotation(book(), client)
        self.assertTrue(result["within"])

        short = retell.annotation(book(), self.FakeClient("коротко"))
        self.assertFalse(short["within"])

    def test_text_is_not_trimmed(self):
        """Обрезанная на полуслове аннотация хуже длинной."""
        client = self.FakeClient("а" * 3000)
        self.assertEqual(len(retell.annotation(book(), client)["text"]), 3000)

    def test_empty_registry_is_refused_before_the_request(self):
        client = self.FakeClient()
        with self.assertRaises(retell.RetellError):
            retell.annotation(Registry(), client)
        self.assertEqual(client.calls, 0)

    def test_empty_answer_is_an_error(self):
        with self.assertRaises(retell.RetellError):
            retell.annotation(book(), self.FakeClient("  "))

    def test_huge_book_is_cut_to_the_start_and_the_end(self):
        """Пятьсот глав в один запрос не влезут, а сюжет виден по краям."""
        registry = Registry()
        for n in range(1, 301):
            registry.add_event(Event(chapter=n, type="шаг", actor="Герой"))
        lines = retell.facts_lines(registry)
        chapters = [l for l in lines if l.startswith("глава ")]
        self.assertLessEqual(len(chapters), retell.CHAPTERS_FOR_ANNOTATION)
        self.assertTrue(any("глава 1:" in l for l in chapters))
        self.assertTrue(any("глава 300:" in l for l in chapters))


class TestExport(unittest.TestCase):
    """Выгрузка карточек и пересказа в файл."""

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def test_markdown_is_written_as_is(self):
        path = docs.save("# Заголовок\n\nТекст.", self.tmp / "карточки.md")
        self.assertIn("# Заголовок", Path(path).read_text(encoding="utf-8"))

    def test_missing_folder_is_created(self):
        path = docs.save("Текст", self.tmp / "нет" / "такой" / "п.md")
        self.assertTrue(Path(path).is_file())

    def test_unknown_format_is_refused(self):
        with self.assertRaises(docs.ExportError):
            docs.save("Текст", self.tmp / "файл.pdf")

    def test_empty_text_is_refused(self):
        with self.assertRaises(docs.ExportError):
            docs.save("   ", self.tmp / "пусто.md")

    def test_docx_gets_real_headings(self):
        try:
            import docx  # noqa: F401
        except ImportError:
            self.skipTest("нет python-docx")

        from docx import Document

        path = docs.save("# Ли Вэй\n\nТип: персонаж", self.tmp / "к.docx")
        document = Document(path)
        styles = [p.style.name for p in document.paragraphs if p.text.strip()]
        self.assertTrue(any(s.startswith("Heading") for s in styles))
        self.assertIn("Ли Вэй", [p.text for p in document.paragraphs])


if __name__ == "__main__":
    unittest.main(verbosity=2)
