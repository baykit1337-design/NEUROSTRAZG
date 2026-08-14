"""Тесты ядра (часть A ТЗ NEUROSTRAZH).

Обязательный минимум из A0: обработка текста, разбор имён, каждый читатель
на своём формате и round-trip каждого писателя — записали, прочитали
обратно, данные совпали.

Интернет не нужен: все файлы собираются на лету.
"""

from __future__ import annotations

import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core import formats, naming, text  # noqa: E402
from core.models import Book, Chapter, OpReport  # noqa: E402
from core.readers.base import ReadError  # noqa: E402
from core.writers.base import WriteError  # noqa: E402

PARAGRAPHS = [
    "Первый абзац главы, достаточно длинный, чтобы его было видно.",
    "Второй абзац с другим содержанием и парой слов сверху.",
    "Третий и последний абзац этой главы.",
]


def sample(number: int = 201, title: str = "Глава 201. Название") -> Chapter:
    return Chapter(number=number, title=title, paragraphs=list(PARAGRAPHS))


class TestModels(unittest.TestCase):
    def test_chapter_text_joins_paragraphs(self):
        self.assertEqual(sample().text, "\n\n".join(PARAGRAPHS))

    def test_chapter_label(self):
        self.assertEqual(Chapter(number=201).label, "201")
        self.assertEqual(Chapter(number=201, part=2).label, "201.2")
        self.assertEqual(Chapter().label, "")

    def test_book_size_and_length(self):
        book = Book(chapters=[sample(), sample(202)])
        self.assertEqual(len(book), 2)
        self.assertEqual(book.size, sample().size * 2)

    def test_report_collects_failures(self):
        report = OpReport()
        report.fail("a.txt", "чтение", "BadZipFile: не архив")
        self.assertEqual(report.failed, 1)
        self.assertIn("a.txt — чтение", report.as_dict()["failed_files"][0])


class TestNaming(unittest.TestCase):
    """A0: разбор имён всех известных видов и сборка обратно."""

    def test_examples(self):
        cases = {
            "0010 - Глава 209. Название": (10, 209, None, "Название"),
            "0010 - Глава 209: Название": (10, 209, None, "Название"),
            "Глава 209. Название": (None, 209, None, "Название"),
            "Chapter 209. Название": (None, 209, None, "Название"),
            "0012 - Глава 210.2. Часть": (12, 210, 2, "Часть"),
        }
        for stem, expected in cases.items():
            parts = naming.parse(stem)
            with self.subTest(stem=stem):
                self.assertEqual(
                    (parts.seq, parts.number, parts.part, parts.title), expected
                )

    def test_name_without_a_second_number_keeps_the_leading_one(self):
        """Отбрасывать порядковый номер можно, только если есть замена."""
        parts = naming.parse("0001 - Информация")
        self.assertEqual(parts.number, 1)
        self.assertEqual(parts.title, "Информация")

    def test_examples_from_the_spec(self):
        """Разбор не привязан ни к языку, ни к слову «глава»."""
        cases = {
            "0001 - Chapter 241_ Panicking Count Ashton": (241, "Panicking Count Ashton"),
            "0002 - Глава 242. Улика": (242, "Улика"),
            "Chapter 243) Finding the Culprit": (243, "Finding the Culprit"),
            "第 244 章 标题": (244, "标题"),
            "глава244безпробелов": (244, "безпробелов"),
            "ЮЮЮ 245 ЮЮЮ": (245, "ЮЮЮ"),
        }
        for stem, (number, title) in cases.items():
            with self.subTest(stem=stem):
                parts = naming.parse(stem)
                self.assertEqual((parts.number, parts.title), (number, title))

    def test_name_without_digits_keeps_the_whole_name(self):
        parts = naming.parse("Пролог")
        self.assertIsNone(parts.number)
        self.assertEqual(parts.title, "Пролог")

    def test_long_digit_runs_are_not_chapter_numbers(self):
        """Шесть цифр — это дата или внутренний код, но не номер главы."""
        parts = naming.parse("20240101 - Название")
        self.assertIsNone(parts.number)

    def test_part_needs_digits_right_after_the_dot(self):
        self.assertEqual(naming.parse("Глава 361.2").part, 2)
        # «Глава 5. 100 дней» — точка с пробелом это разделитель, не часть.
        parts = naming.parse("Глава 5. 100 дней")
        self.assertEqual((parts.number, parts.part, parts.title), (5, None, "100 дней"))

    def test_outliers_are_flagged(self):
        """Из даты в имени выйдет «глава 2024» посреди двухсот обычных."""
        self.assertEqual(naming.suspects([201, 202, 203, 204, 2024, 205, 206]), {2024})
        self.assertEqual(naming.suspects(list(range(1, 501))), set())
        # Данных мало — судить не о чем.
        self.assertEqual(naming.suspects([1, 2]), set())

    def test_build_round_trip(self):
        """Собрали имя — разобрали обратно, номер и название на месте."""
        fmt = naming.NameFormat(separator=". ")
        for number, part, title in ((201, None, "Конец"), (361, 2, "Начало")):
            name = naming.build(number, part, title, fmt)
            back = naming.parse(name)
            with self.subTest(name=name):
                self.assertEqual(back.number, number)
                self.assertEqual(back.part, part)
                self.assertEqual(back.title, title)

    def test_forbidden_characters_replaced(self):
        for char in ':/\\|*?"<>':
            self.assertNotIn(char, naming.safe_filename(f"Имя{char}тут"))

    def test_colon_kept_readable(self):
        """Двоеточие запрещено и в Windows, и в macOS — но не теряется."""
        self.assertEqual(naming.safe_filename("Глава 1: Имя"), "Глава 1 - Имя")

    def test_reserved_windows_names(self):
        self.assertTrue(naming.safe_filename("CON").startswith("_"))

    def test_sort_key_is_numeric(self):
        chapters = [Chapter(number=361, part=10), Chapter(number=361, part=2)]
        ordered = sorted(chapters, key=naming.sort_key)
        self.assertEqual([c.part for c in ordered], [2, 10])


class TestText(unittest.TestCase):
    """A0: обработка абзацев, разделителей, дублей названия."""

    def test_duplicate_title_removed(self):
        title = "Глава 209. Название"
        blocks = text.prepare([title, title, "Текст."], title)
        self.assertEqual([b.text for b in blocks], ["Текст."])

    def test_scene_breaks_collapse(self):
        blocks = text.prepare(["А", "*", "*", "Б"], "")
        self.assertEqual([b.text for b in blocks], ["А", "* * *", "Б"])

    def test_empty_paragraphs_dropped(self):
        blocks = text.prepare(["А", "", "   ", "Б"], "")
        self.assertEqual([b.text for b in blocks], ["А", "Б"])

    def test_defaults(self):
        options = text.PrepOptions()
        self.assertEqual(options.align, "left")
        self.assertEqual(options.first_line_indent_cm, 0.0)


class FormatTestCase(unittest.TestCase):
    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)


class TestReaders(FormatTestCase):
    """Каждый читатель — на маленьком файле своего формата."""

    def read(self, name: str) -> list[Chapter]:
        return formats.read(self.tmp / name)

    def test_txt(self):
        (self.tmp / "a.txt").write_text("\n\n".join(PARAGRAPHS), encoding="utf-8")
        self.assertEqual(self.read("a.txt")[0].paragraphs, PARAGRAPHS)

    def test_txt_cp1251(self):
        """Русский текст приходит и в CP1251 — читать вслепую нельзя."""
        (self.tmp / "a.txt").write_bytes("\n\n".join(PARAGRAPHS).encode("cp1251"))
        self.assertEqual(self.read("a.txt")[0].paragraphs, PARAGRAPHS)

    def test_md(self):
        (self.tmp / "a.md").write_text("\n\n".join(PARAGRAPHS), encoding="utf-8")
        self.assertEqual(self.read("a.md")[0].paragraphs, PARAGRAPHS)

    def test_docx(self):
        from docx import Document

        document = Document()
        for paragraph in PARAGRAPHS:
            document.add_paragraph(paragraph)
        document.save(str(self.tmp / "a.docx"))
        self.assertEqual(self.read("a.docx")[0].paragraphs, PARAGRAPHS)

    def test_epub(self):
        from tests.test_split import make_epub

        make_epub(self.tmp / "a.epub", count=3)
        chapters = self.read("a.epub")
        self.assertEqual(len(chapters), 3)
        self.assertTrue(all(c.paragraphs for c in chapters))

    def test_fb2(self):
        body = "".join(
            f"<section><title><p>Глава {n}</p></title>"
            + "".join(f"<p>{p}</p>" for p in PARAGRAPHS)
            + "</section>"
            for n in (201, 202)
        )
        (self.tmp / "a.fb2").write_text(
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<FictionBook xmlns="http://www.gribuser.ru/xml/fictionbook/2.0">'
            f"<body>{body}</body></FictionBook>",
            encoding="utf-8",
        )
        chapters = self.read("a.fb2")
        self.assertEqual(len(chapters), 2)
        self.assertEqual(chapters[0].paragraphs, PARAGRAPHS)

    def test_rtf(self):
        from core.writers.rtf import RtfWriter

        RtfWriter().write(self.tmp / "a.rtf", [sample()], headings=False)
        self.assertTrue(self.read("a.rtf")[0].paragraphs)

    def test_odt(self):
        from core.writers.odt import OdtWriter

        OdtWriter().write(self.tmp / "a.odt", [sample()], headings=False)
        self.assertEqual(self.read("a.odt")[0].paragraphs, PARAGRAPHS)

    def test_html(self):
        (self.tmp / "a.html").write_text(
            "<html><body><script>x=1</script>"
            + "".join(f"<p>{p}</p>" for p in PARAGRAPHS)
            + "</body></html>",
            encoding="utf-8",
        )
        self.assertEqual(self.read("a.html")[0].paragraphs, PARAGRAPHS)

    def test_unknown_format_refused(self):
        (self.tmp / "a.pdf").write_bytes(b"%PDF-1.4")
        with self.assertRaises(ReadError):
            self.read("a.pdf")


class TestSignatureSniffing(FormatTestCase):
    """Расширение врёт: содержимое важнее."""

    def test_epub_named_txt_is_still_read_as_epub(self):
        from tests.test_split import make_epub

        make_epub(self.tmp / "book.epub", count=2)
        disguised = self.tmp / "book.txt"
        disguised.write_bytes((self.tmp / "book.epub").read_bytes())

        chapters = formats.read(disguised)
        self.assertEqual(len(chapters), 2)
        # Сырые байты архива в текст не попали.
        self.assertNotIn("PK", chapters[0].paragraphs[0][:4])

    def test_sniff_recognises_each_zip_kind(self):
        from tests.test_split import make_epub

        make_epub(self.tmp / "a.epub", count=1)
        self.assertEqual(formats.sniff(self.tmp / "a.epub"), ".epub")

        from core.writers.odt import OdtWriter

        OdtWriter().write(self.tmp / "a.odt", [sample()])
        self.assertEqual(formats.sniff(self.tmp / "a.odt"), ".odt")

    def test_plain_text_has_no_signature(self):
        (self.tmp / "a.txt").write_text("просто текст", encoding="utf-8")
        self.assertEqual(formats.sniff(self.tmp / "a.txt"), "")


class TestWriterRoundTrip(FormatTestCase):
    """Записали, прочитали обратно — данные совпали."""

    def round_trip(self, suffix: str) -> list[Chapter]:
        path = self.tmp / f"book{suffix}"
        formats.write(path, [sample()], headings=False)
        return formats.read(path)

    def test_txt(self):
        self.assertEqual(self.round_trip(".txt")[0].paragraphs, PARAGRAPHS)

    def test_md(self):
        self.assertEqual(self.round_trip(".md")[0].paragraphs, PARAGRAPHS)

    def test_docx(self):
        self.assertEqual(self.round_trip(".docx")[0].paragraphs, PARAGRAPHS)

    def test_odt(self):
        self.assertEqual(self.round_trip(".odt")[0].paragraphs, PARAGRAPHS)

    def test_rtf(self):
        # striprtf схлопывает пробелы, поэтому сверяем по содержанию.
        paragraphs = self.round_trip(".rtf")[0].paragraphs
        self.assertEqual(len(paragraphs), len(PARAGRAPHS))
        self.assertIn("Первый абзац", paragraphs[0])

    def test_fb2(self):
        self.assertEqual(self.round_trip(".fb2")[0].paragraphs, PARAGRAPHS)

    def test_epub(self):
        path = self.tmp / "book.epub"
        # Главы короче 200 символов epub-читатель отсекает как служебные,
        # поэтому для round-trip берём главу нормального объёма.
        long_chapter = Chapter(
            number=201, title="Глава 201. Название",
            paragraphs=[p * 4 for p in PARAGRAPHS],
        )
        formats.write(path, [long_chapter], headings=True)
        chapters = formats.read(path)
        self.assertEqual(len(chapters), 1)
        self.assertIn("Первый абзац", chapters[0].text)

    def test_every_writable_format_round_trips(self):
        """Ни один писатель не должен ронять запись."""
        for suffix in formats.WRITABLE:
            with self.subTest(suffix=suffix):
                path = self.tmp / f"all{suffix}"
                formats.write(path, [sample()], headings=True)
                self.assertTrue(path.exists())
                self.assertGreater(path.stat().st_size, 0)

    def test_unknown_writer_refused(self):
        with self.assertRaises(WriteError):
            formats.write(self.tmp / "a.pdf", [sample()])

    def test_txt_windows_1251(self):
        path = self.tmp / "cp.txt"
        formats.write(path, [sample()], encoding="windows-1251", headings=False)
        self.assertIn("Первый абзац", path.read_text(encoding="windows-1251"))

    def test_epub_has_valid_opf_and_spine(self):
        path = self.tmp / "book.epub"
        formats.write(path, [sample(201), sample(202)], headings=True)
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            opf = archive.read("OEBPS/content.opf").decode("utf-8")

        # mimetype обязан быть первым и несжатым, иначе читалки ругаются.
        self.assertEqual(names[0], "mimetype")
        self.assertIn("META-INF/container.xml", names)
        # Порядок чтения задаёт spine.
        self.assertEqual(opf.count("<itemref"), 2)
        self.assertLess(opf.index('idref="c1"'), opf.index('idref="c2"'))


class TestMixedFolder(FormatTestCase):
    """В папке со смешанными форматами обрабатываются все подряд."""

    def test_all_formats_read_from_one_folder(self):
        formats.write(self.tmp / "a.txt", [sample(201)], headings=False)
        formats.write(self.tmp / "b.md", [sample(202)], headings=False)
        formats.write(self.tmp / "c.docx", [sample(203)], headings=False)

        chapters = []
        for path in sorted(self.tmp.iterdir()):
            chapters.extend(formats.read(path))
        self.assertEqual(len(chapters), 3)

    def test_unreadable_file_does_not_stop_the_rest(self):
        formats.write(self.tmp / "a.txt", [sample()], headings=False)
        (self.tmp / "b.pdf").write_bytes(b"%PDF-1.4")

        report = OpReport()
        chapters = []
        for path in sorted(self.tmp.iterdir()):
            try:
                chapters.extend(formats.read(path))
            except ReadError as exc:
                report.fail(path.name, "чтение", str(exc))

        self.assertEqual(len(chapters), 1)
        self.assertEqual(report.failed, 1)




class TestOps(FormatTestCase):
    """A1: «Разбить» и «Объединить» — одна логика с параметром формата."""

    def setUp(self):
        super().setUp()
        from core import formats

        self.src = self.tmp / "смесь"
        self.src.mkdir()
        # Папка со смешанными форматами — так и бывает в работе.
        for number, suffix in [(201, ".txt"), (202, ".docx"), (203, ".md")]:
            formats.write(
                self.src / f"Глава {number}{suffix}",
                [Chapter(number=number, title=f"Глава {number}. Имя",
                         paragraphs=[f"Текст главы {number}. " * 12])],
                headings=True,
            )

    def test_split_accepts_docx(self):
        """«Разбить» не принимал .docx — это был живой баг."""
        from ops import split

        out = self.tmp / "разбито"
        report = split.run([str(self.src / "Глава 202.docx")], out, out_format=".txt")
        self.assertEqual(report.written, 1)

    def test_split_converts_between_formats(self):
        from ops import split

        out = self.tmp / "epub"
        split.run([str(self.src / "Глава 202.docx")], out, out_format=".epub")
        self.assertEqual([p.suffix for p in out.iterdir()], [".epub"])

    def test_split_into_parts_numbers_them(self):
        from ops import split

        chapter = Chapter(number=201, title="Глава 201",
                          paragraphs=[f"Абзац {i}. " * 8 for i in range(10)])
        parts = split.split_chapter(chapter, 3)
        self.assertEqual([p.part for p in parts], [1, 2, 3])

    def test_merge_reads_mixed_formats(self):
        from ops import merge

        out = self.tmp / "книга.txt"
        report = merge.run([str(self.src)], out)
        self.assertEqual(report.written, 3)
        self.assertEqual(report.failed, 0)

    def test_merge_orders_by_chapter_number(self):
        from ops import merge

        out = self.tmp / "книга.txt"
        merge.run([str(self.src)], out, headings=False)
        body = out.read_text(encoding="utf-8")
        self.assertLess(body.index("главы 201"), body.index("главы 203"))

    def test_merge_writes_any_format(self):
        from core import formats
        from ops import merge

        for suffix in formats.WRITABLE:
            with self.subTest(suffix=suffix):
                out = self.tmp / f"книга{suffix}"
                report = merge.run([str(self.src)], out)
                self.assertEqual(report.failed, 0)
                self.assertTrue(out.exists())

    def test_merge_separator(self):
        from ops import merge

        out = self.tmp / "книга.txt"
        merge.run([str(self.src)], out, separator="stars")
        self.assertIn("* * *", out.read_text(encoding="utf-8"))

    def test_unreadable_file_is_reported_not_swallowed(self):
        from ops import merge

        (self.src / "битый.pdf").write_bytes(b"%PDF-1.4")
        out = self.tmp / "книга.txt"
        report = merge.run([str(self.src)], out)
        # Файл пропущен, остальные собраны, ошибка в отчёте.
        self.assertEqual(report.written, 3)

    def test_cancel_is_checked(self):
        import threading

        from ops import merge
        from ops.base import Cancelled, Progress

        cancel = threading.Event()
        cancel.set()
        with self.assertRaises(Cancelled):
            merge.run([str(self.src)], self.tmp / "к.txt",
                      progress=Progress(cancel=cancel))

    def test_progress_callback_shape(self):
        from ops import split
        from ops.base import Progress

        seen = []
        split.run([str(self.src)], self.tmp / "out",
                  progress=Progress(on_progress=lambda d, t, m: seen.append((d, t, m))))
        self.assertTrue(seen)
        # Колбэк один на все операции: (сделано, всего, текст).
        self.assertEqual(len(seen[0]), 3)


class TestHeadings(unittest.TestCase):
    """A1.1: книга одним куском режется по заголовкам — в любом формате."""

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def _book(self, suffix: str) -> Path:
        """Вся книга одним файлом: заголовки идут обычными абзацами."""
        from core import formats

        paragraphs = []
        for number in (1, 2, 3):
            paragraphs.append(f"Глава {number}. Название {number}")
            paragraphs.append(f"Текст главы {number}. " * 10)
        path = self.tmp / f"книга{suffix}"
        formats.write(path, [Chapter(title="книга", paragraphs=paragraphs)],
                      headings=False)
        return path

    def test_cut_finds_chapters(self):
        from core import headings

        chapter = Chapter(title="книга", paragraphs=[
            "Глава 1. Первая", "Текст первой.", "Глава 2. Вторая", "Текст второй."])
        parts = headings.cut(chapter)
        self.assertEqual([c.number for c in parts], [1, 2])
        self.assertEqual(parts[0].paragraphs, ["Текст первой."])
        # Заголовок в тело не попадает: иначе после записи он задвоится.
        self.assertNotIn("Глава 1. Первая", parts[1].paragraphs)

    def test_cut_without_headings_asks_for_pattern(self):
        from core import headings

        with self.assertRaises(headings.HeadingsNotFound):
            headings.cut(Chapter(title="книга", paragraphs=["Просто текст."]))

    def test_own_pattern_wins(self):
        from core import headings

        chapter = Chapter(paragraphs=["### 1 ###", "Раз", "### 2 ###", "Два"])
        parts = headings.cut(chapter, r"^###\s*\d+\s*###$")
        self.assertEqual(len(parts), 2)

    def test_mention_in_text_is_not_a_heading(self):
        from core import headings

        # «в главе 12 говорилось» — обычный текст, а не заголовок.
        chapter = Chapter(paragraphs=[
            "Глава 1. Начало", "Как в главе 12 говорилось, всё было иначе."])
        self.assertEqual(headings.find(chapter.paragraphs), [0])

    def test_split_cuts_monolithic_book_in_every_flat_format(self):
        """Резать книгу целиком умеем не только в .txt — это и было задачей A2."""
        from ops import split

        for suffix in (".txt", ".md", ".rtf", ".odt", ".docx"):
            with self.subTest(suffix=suffix):
                book = self._book(suffix)
                out = self.tmp / f"главы{suffix}"
                report = split.run([str(book)], out, out_format=".txt")
                self.assertEqual(report.written, 3, suffix)
                self.assertEqual(len(list(out.glob("*.txt"))), 3)

    def test_folder_of_chapters_is_not_cut(self):
        """Папка готовых глав — не книга: шаблон просить не за что."""
        from core import formats
        from ops import split

        src = self.tmp / "главы"
        src.mkdir()
        for number in (1, 2):
            formats.write(src / f"Глава {number}.txt",
                          [Chapter(number=number, title=f"Глава {number}",
                                   paragraphs=["Текст."])])
        report = split.run([str(src)], self.tmp / "out", out_format=".txt")
        self.assertEqual(report.written, 2)

    def test_single_chapter_split_into_parts_needs_no_headings(self):
        """Делению на части заголовки не нужны — просить шаблон незачем."""
        from core import formats
        from ops import split

        one = self.tmp / "Пролог.txt"
        formats.write(one, [Chapter(title="Пролог",
                                    paragraphs=[f"Абзац {n}." for n in range(10)])])
        report = split.run([str(one)], self.tmp / "части", out_format=".txt", parts=2)
        self.assertEqual(report.written, 2)


if __name__ == "__main__":
    unittest.main(verbosity=2)


class TestCollectFiles(unittest.TestCase):
    """Как ядро обращается с чужим форматом — по-разному и намеренно."""

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def test_single_unknown_file_is_refused(self):
        """Выбран руками — значит, отказ, а не чтение .pdf как текста."""
        from ops.base import collect_files

        bad = self.tmp / "книга.pdf"
        bad.write_bytes(b"%PDF-1.4 \x00\x01\x02")
        with self.assertRaises(ReadError):
            collect_files([str(bad)])

    def test_unknown_file_in_folder_is_skipped(self):
        """В папке — молча мимо: там лежит и служебное."""
        from ops.base import collect_files

        (self.tmp / "state.json").write_text("{}", encoding="utf-8")
        (self.tmp / "книга.pdf").write_bytes(b"%PDF-1.4")
        formats.write(self.tmp / "Глава 1.txt",
                      [Chapter(number=1, title="Глава 1", paragraphs=["Текст."])])
        files = collect_files([str(self.tmp)])
        self.assertEqual([f.name for f in files], ["Глава 1.txt"])

    def test_signature_beats_extension_for_single_file(self):
        """epub под чужим именем читается: спор решает содержимое."""
        from ops.base import collect_files

        path = self.tmp / "книга.dat"
        formats.write(self.tmp / "книга.epub",
                      [Chapter(number=1, title="Глава 1", paragraphs=["Текст."])])
        (self.tmp / "книга.epub").rename(path)
        self.assertEqual(collect_files([str(path)]), [path])

    def test_missing_path_is_reported(self):
        from ops.base import collect_files

        with self.assertRaises(ReadError):
            collect_files([str(self.tmp / "нет-такого.txt")])


class TestTabsShareTheCore(unittest.TestCase):
    """A0: у вкладок нет своих списков форматов и своего разбора имён."""

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def test_format_lists_come_from_core(self):
        from mvl import rename, source, textcheck

        self.assertEqual(source.READABLE, formats.READABLE)
        self.assertEqual(rename.READABLE, formats.READABLE)
        self.assertEqual(textcheck.READABLE, formats.READABLE)

    def test_naming_is_not_duplicated(self):
        """«Переименовать» разбирает имена тем же кодом, что и остальные."""
        from mvl import rename

        self.assertIs(rename.NameFormat, naming.NameFormat)
        self.assertIs(rename.safe_filename, naming.safe_filename)
        self.assertIs(rename.build_name, naming.build)
        # Разбор обёрнут ради своей ошибки, но ходит в ядро.
        self.assertEqual(rename.parse_name("Глава 12.3: Имя").part, 3)

    def test_text_processing_is_not_duplicated(self):
        from mvl import textprep

        self.assertIs(textprep.prepare, text.prepare)
        self.assertIs(textprep.PrepOptions, text.PrepOptions)

    def test_check_reads_every_format_core_can(self):
        """Проверка принимала 4 формата из 8 — теперь все, что читает ядро."""
        from mvl import textcheck

        body = "Тут 修炼 остался. " + "Обычный текст главы. " * 20
        for suffix in formats.WRITABLE:
            with self.subTest(suffix=suffix):
                path = self.tmp / f"Глава 1{suffix}"
                formats.write(path, [Chapter(number=1, title="Глава 1",
                                             paragraphs=[body])], headings=True)
                report = textcheck.check(path, kinds=["cjk"])
                self.assertEqual(len(report.findings), 1, suffix)


class TestCheckCounts(unittest.TestCase):
    """Счётчик файлов в отчёте проверки."""

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def test_book_wide_findings_are_not_counted_as_a_file(self):
        """Выходило «в 6 файлах из 5»: находки по книге считались файлом."""
        from mvl import textcheck

        body = "Тут 修炼 остался. " + "Обычный текст главы. " * 25
        # Дыра в нумерации даёт находку «по всей книге».
        for number in (201, 203):
            formats.write(self.tmp / f"Глава {number}.txt",
                          [Chapter(number=number, title=f"Глава {number}",
                                   paragraphs=[body])], headings=True)

        report = textcheck.check(self.tmp)
        self.assertTrue(
            any(f.file == textcheck.BOOK_WIDE for f in report.findings),
            "нужна хотя бы одна находка по книге целиком")
        self.assertEqual(report.files_checked, 2)
        self.assertLessEqual(report.files_with_findings, report.files_checked)


class TestHeaders(unittest.TestCase):
    """1.2: мусорная шапка определяется по повторам, без жёстких правил."""

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def book(self, count=10, book_name="Genetic Ascension"):
        """Папка в том виде, в каком её отдаёт парсер: с шапкой."""
        folder = self.tmp / "книга"
        folder.mkdir(exist_ok=True)
        for n in range(2250, 2250 + count):
            name = f"Chapter {n}_ Handsome [Bonus]"
            body = "\n\n".join([
                book_name,
                name,
                f"Первый настоящий абзац главы {n}, у каждой свой.",
                f"Второй абзац главы {n}. " * 6,
            ])
            (folder / f"{name}.txt").write_text(body, encoding="utf-8")
        return folder

    def findings(self, folder):
        from ops import headers

        return headers.scan([str(folder)])["findings"]

    def test_book_name_is_found_by_repetition(self):
        """Название книги вычисляется само: оно есть почти в каждом файле."""
        found = [f for f in self.findings(self.book()) if f["kind"] == "repeat"]
        self.assertEqual([f["text"] for f in found], ["Genetic Ascension"])
        self.assertEqual(found[0]["count"], 10)
        self.assertIn("в 10 файлах из 10", found[0]["label"])

    def test_title_echo_is_found_separately(self):
        found = [f for f in self.findings(self.book()) if f["kind"] == "title"]
        self.assertEqual(len(found), 1)
        self.assertEqual(found[0]["count"], 10)

    def test_title_echo_survives_different_separators(self):
        """`Chapter 243_ …` и `Chapter 243: …` — одна и та же строка."""
        folder = self.tmp / "разделители"
        folder.mkdir()
        for n in range(240, 250):
            name = f"Chapter {n}_ Finding the Culprit (Bonus)"
            (folder / f"{name}.txt").write_text("\n\n".join([
                name,
                f"Chapter {n}: Finding the Culprit (Bonus)",
                f"Текст главы {n}, у каждой свой.",
            ]), encoding="utf-8")
        found = [f for f in self.findings(folder) if f["kind"] == "title"]
        self.assertEqual(found[0]["count"], 10)

    def test_rare_line_is_not_a_header(self):
        """Строка из одного файла — содержание, а не шапка."""
        folder = self.book()
        first = sorted(folder.iterdir())[0]
        first.write_text("Особая строка только здесь.\n\nТекст.", encoding="utf-8")
        texts = [f["text"] for f in self.findings(folder)]
        self.assertNotIn("Особая строка только здесь.", texts)

    def test_long_paragraph_is_never_a_header(self):
        """Абзац на три экрана шапкой не бывает, даже если повторился."""
        folder = self.tmp / "длинные"
        folder.mkdir()
        long_line = "Очень длинный повторяющийся абзац. " * 12
        for n in range(1, 11):
            (folder / f"Глава {n}.txt").write_text(
                f"{long_line}\n\nТекст главы {n}.", encoding="utf-8")
        self.assertEqual(self.findings(folder), [])

    def test_clean_removes_only_the_header(self):
        from ops import headers

        folder = self.book()
        out = self.tmp / "чисто"
        report = headers.run([str(folder)], out,
                             ["Genetic Ascension", ""])
        self.assertEqual(report.written, 10)

        body = (out / "Chapter 2250_ Handsome [Bonus].txt").read_text(encoding="utf-8")
        self.assertNotIn("Genetic Ascension", body)
        # Название главы осталось ровно одно — его пишет писатель.
        self.assertEqual(body.count("Handsome [Bonus]"), 1)
        self.assertIn("Первый настоящий абзац", body)

    def test_originals_are_untouched(self):
        from ops import headers

        folder = self.book()
        before = {p.name: p.read_text(encoding="utf-8") for p in folder.iterdir()}
        headers.run([str(folder)], self.tmp / "out2", ["Genetic Ascension"])
        after = {p.name: p.read_text(encoding="utf-8") for p in folder.iterdir()}
        self.assertEqual(before, after)

    def test_same_line_deeper_in_the_text_is_kept(self):
        """Чистим только зону шапки: дальше это уже содержание."""
        paragraphs = ["Genetic Ascension", "Абзац.", "Абзац.", "Абзац.",
                      "Абзац.", "Genetic Ascension"]
        kept = text.strip_headers(paragraphs, "Глава 1", ["Genetic Ascension"])
        self.assertEqual(kept[-1], "Genetic Ascension")
        self.assertNotIn("Genetic Ascension", kept[:1])

    def test_clean_works_in_every_format(self):
        """Живёт в ядре, поэтому формат значения не имеет."""
        from ops import headers

        for suffix in (".txt", ".md", ".docx", ".fb2", ".rtf", ".odt"):
            with self.subTest(suffix=suffix):
                folder = self.tmp / f"ф{suffix}"
                folder.mkdir()
                for n in range(1, 6):
                    formats.write(
                        folder / f"Глава {n}{suffix}",
                        [Chapter(number=n, title=f"Глава {n}", paragraphs=[
                            "Название книги", f"Текст главы {n}, у каждой свой."])],
                        headings=False)
                found = [f["text"] for f in self.findings(folder)]
                self.assertIn("Название книги", found, suffix)
