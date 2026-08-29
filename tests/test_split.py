"""Тесты вкладки «Разбить книгу».

Тестовые epub собираются на лету — так проверка не зависит от того, какие
файлы лежат рядом.
"""

from __future__ import annotations

import sys
import threading
import unittest

from core import headings as corehead
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from mvl import booksplit  # noqa: E402
from mvl.word import Style, split_paragraphs  # noqa: E402

LONG = "Строка текста этой главы, достаточно длинная, чтобы пройти отсечку. " * 5


def make_epub(path: Path, count: int = 3, with_cover: bool = True) -> Path:
    """Минимальный, но настоящий epub: container.xml, OPF, manifest, spine."""
    items, refs, files = [], [], {}

    if with_cover:
        # Служебная страница: короткая, должна отсечься по длине.
        files["OEBPS/cover.xhtml"] = "<html><body><p>Обложка</p></body></html>"
        items.append('<item id="cover" href="cover.xhtml" media-type="application/xhtml+xml"/>')
        refs.append('<itemref idref="cover"/>')

    for n in range(1, count + 1):
        body = "".join(f"<p>Абзац {i}. {LONG}</p>" for i in range(1, 4))
        files[f"OEBPS/ch{n}.xhtml"] = (
            f"<html><body><h1>Глава {n}: Название</h1>{body}"
            f"<p>*</p><p>После разделителя. {LONG}</p></body></html>"
        )
        items.append(f'<item id="c{n}" href="ch{n}.xhtml" media-type="application/xhtml+xml"/>')
        refs.append(f'<itemref idref="c{n}"/>')

    with zipfile.ZipFile(path, "w") as archive:
        archive.writestr(
            "META-INF/container.xml",
            '<?xml version="1.0"?><container xmlns="urn:oasis:names:tc:opendocument:xmlns:container">'
            '<rootfiles><rootfile full-path="OEBPS/book.opf"/></rootfiles></container>',
        )
        for name, content in files.items():
            archive.writestr(name, content)
        archive.writestr(
            "OEBPS/book.opf",
            f'<?xml version="1.0"?><package xmlns="http://www.idpf.org/2007/opf">'
            f'<manifest>{"".join(items)}</manifest><spine>{"".join(refs)}</spine></package>',
        )
    return path


def make_txt(path: Path, count: int = 3) -> Path:
    parts = ["Предисловие, которое не должно стать главой.\n"]
    for n in range(1, count + 1):
        parts.append(f"Глава {n}. Название\n\n{LONG}\n\n*\n\nПосле разделителя.\n")
    path.write_text("\n".join(parts), encoding="utf-8")
    return path


class SplitTestCase(unittest.TestCase):
    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)
        self.epub = make_epub(self.tmp / "book.epub")
        self.txt = make_txt(self.tmp / "book.txt")


class TestReadChapters(SplitTestCase):
    def test_epub_chapters_in_reading_order(self):
        chapters = booksplit.read_chapters(self.epub)
        self.assertEqual([c.title for c in chapters],
                         ["Глава 1: Название", "Глава 2: Название", "Глава 3: Название"])

    def test_cover_page_filtered_out(self):
        # Обложка короче 200 символов — в главы не попадает.
        self.assertEqual(len(booksplit.read_chapters(self.epub)), 3)

    def test_scene_separator_kept(self):
        text = booksplit.read_chapters(self.epub)[0].text
        self.assertIn("\n\n*\n\n", text)

    def test_txt_chapters(self):
        chapters = booksplit.read_chapters(self.txt)
        self.assertEqual([c.title for c in chapters],
                         ["Глава 1. Название", "Глава 2. Название", "Глава 3. Название"])

    def test_txt_scene_separator_kept(self):
        self.assertIn("\n*\n", booksplit.read_chapters(self.txt)[0].text)

    def test_custom_pattern(self):
        path = self.tmp / "custom.txt"
        path.write_text(f"### 1 Первая\n\n{LONG}\n\n### 2 Вторая\n\n{LONG}\n", encoding="utf-8")
        chapters = booksplit.read_chapters(path, pattern=r"^###\s+\d+.*$")
        self.assertEqual(len(chapters), 2)

    def test_missing_headings_raises_with_pattern(self):
        path = self.tmp / "flat.txt"
        path.write_text("Текст без единого заголовка главы.\n", encoding="utf-8")
        with self.assertRaises(booksplit.HeadingsNotFound) as ctx:
            booksplit.read_chapters(path)
        self.assertEqual(ctx.exception.pattern, booksplit.DEFAULT_PATTERN)

    def test_missing_file_raises_split_error(self):
        with self.assertRaises(booksplit.SplitError):
            booksplit.read_chapters(self.tmp / "nope.epub")

    def test_unknown_extension_rejected(self):
        path = self.tmp / "book.pdf"
        path.write_text("x", encoding="utf-8")
        with self.assertRaises(booksplit.SplitError):
            booksplit.read_chapters(path)

    def test_broken_epub_does_not_raise_systemexit(self):
        """SystemExit из split_book не должен всплывать в сервер."""
        path = self.tmp / "broken.epub"
        with zipfile.ZipFile(path, "w") as archive:
            archive.writestr("nothing.txt", "не epub вовсе")
        with self.assertRaises(booksplit.SplitError):
            booksplit.read_chapters(path)


class TestPreview(SplitTestCase):
    def test_preview_counts_and_first_titles(self):
        preview = booksplit.preview(self.epub)
        self.assertEqual(preview.total, 3)
        self.assertEqual(preview.titles[0], "Глава 1: Название")
        self.assertEqual(preview.kind, "epub")

    def test_preview_caps_titles_at_five(self):
        big = make_epub(self.tmp / "big.epub", count=9)
        self.assertEqual(len(booksplit.preview(big).titles), 5)

    def test_preview_writes_nothing(self):
        before = set(self.tmp.iterdir())
        booksplit.preview(self.epub)
        self.assertEqual(set(self.tmp.iterdir()), before)


class TestWriteChapters(SplitTestCase):
    def test_txt_filenames_zero_padded(self):
        out = self.tmp / "out"
        booksplit.split_book_to_dir(self.epub, out)
        self.assertEqual(
            sorted(p.name for p in out.glob("*.txt")),
            ["0001 - Глава 1 Название.txt",
             "0002 - Глава 2 Название.txt",
             "0003 - Глава 3 Название.txt"],
        )

    def test_width_grows_with_chapter_count(self):
        self.assertEqual(booksplit.name_width(3), 4)
        self.assertEqual(booksplit.name_width(9999), 4)
        self.assertEqual(booksplit.name_width(10000), 5)

    def test_file_starts_with_title_then_blank_line(self):
        out = self.tmp / "out"
        booksplit.split_book_to_dir(self.epub, out)
        lines = (out / "0001 - Глава 1 Название.txt").read_text(encoding="utf-8").split("\n")
        self.assertEqual(lines[0], "Глава 1: Название")
        self.assertEqual(lines[1], "")

    def test_scene_separator_survives_the_write(self):
        """Разделитель сохраняется, но приводится к единому виду (ТЗ v4)."""
        out = self.tmp / "out"
        booksplit.split_book_to_dir(self.epub, out)
        body = (out / "0001 - Глава 1 Название.txt").read_text(encoding="utf-8")
        self.assertIn("\n\n* * *\n\n", body)

    def test_scene_separator_form_is_configurable(self):
        from mvl.textprep import PrepOptions

        out = self.tmp / "dashes"
        booksplit.split_book_to_dir(
            self.epub, out, prep=PrepOptions(scene_style="dashes")
        )
        body = (out / "0001 - Глава 1 Название.txt").read_text(encoding="utf-8")
        self.assertIn("— — —", body)

    def test_report_counts(self):
        report = booksplit.split_book_to_dir(self.epub, self.tmp / "out")
        self.assertEqual((report.written, report.failed, report.total), (3, 0, 3))

    def test_docx_output(self):
        out = self.tmp / "docx"
        report = booksplit.split_book_to_dir(self.epub, out, fmt=booksplit.FORMAT_DOCX)
        self.assertEqual(report.written, 3)
        self.assertEqual(len(list(out.glob("*.docx"))), 3)

    def test_docx_has_heading_and_paragraphs(self):
        from docx import Document

        out = self.tmp / "docx"
        booksplit.split_book_to_dir(self.epub, out, fmt=booksplit.FORMAT_DOCX)
        document = Document(str(sorted(out.glob("*.docx"))[0]))
        texts = [p.text for p in document.paragraphs if p.text.strip()]
        self.assertEqual(texts[0], "Глава 1: Название")
        self.assertIn("* * *", texts)  # разделитель сцен на месте, в едином виде

    def test_docx_applies_style(self):
        from docx import Document

        out = self.tmp / "styled"
        booksplit.split_book_to_dir(
            self.epub, out, fmt=booksplit.FORMAT_DOCX, style=Style(font="Arial", size=14)
        )
        document = Document(str(sorted(out.glob("*.docx"))[0]))
        self.assertEqual(document.styles["Normal"].font.name, "Arial")

    def test_bad_chapter_does_not_stop_the_rest(self):
        """Ошибка на одной главе — пропуск с записью в отчёт, остальные пишутся."""
        out = self.tmp / "partial"
        out.mkdir()
        chapters = booksplit.read_chapters(self.epub)
        chapters[1] = booksplit.Chapter(title="Плохая", text="текст")
        # Занимаем имя второй главы папкой — запись в неё не пройдёт.
        (out / "0002 - Плохая.txt").mkdir()

        report = booksplit.write_chapters(chapters, out)
        self.assertEqual(report.written, 2)
        self.assertEqual(report.failed, 1)
        self.assertTrue(report.failed_files[0].startswith("0002 - Плохая.txt"))
        # Первая и третья главы записались, несмотря на сбой второй.
        self.assertTrue((out / "0001 - Глава 1 Название.txt").is_file())
        self.assertTrue((out / "0003 - Глава 3 Название.txt").is_file())

    def test_cancel_stops_writing(self):
        cancel = threading.Event()
        cancel.set()
        with self.assertRaises(booksplit.Cancelled):
            booksplit.split_book_to_dir(self.epub, self.tmp / "cancelled", cancel=cancel)

    def test_progress_reported(self):
        seen = []
        booksplit.split_book_to_dir(
            self.epub, self.tmp / "out", on_progress=lambda d, t: seen.append((d, t))
        )
        self.assertEqual(seen, [(1, 3), (2, 3), (3, 3)])

    def test_forbidden_characters_stripped_from_names(self):
        chapters = [booksplit.Chapter(title='Гл 1: A/B "х" <тег>?', text="текст")]
        booksplit.write_chapters(chapters, self.tmp / "names")
        name = next((self.tmp / "names").glob("*.txt")).name
        for char in '\\/:*?"<>|':
            self.assertNotIn(char, name)


class TestWordHelpers(unittest.TestCase):
    def test_split_paragraphs(self):
        self.assertEqual(split_paragraphs("a\n\nb\n\n\nc"), ["a", "b", "c"])

    def test_split_paragraphs_empty(self):
        self.assertEqual(split_paragraphs(""), [])

    def test_style_defaults_from_spec(self):
        style = Style()
        self.assertEqual(style.font, "Times New Roman")
        self.assertEqual(style.size, 12)
        self.assertEqual(style.line_spacing, 1.5)
        # По ТЗ v4 красной строки по умолчанию нет.
        self.assertEqual(style.first_line_indent_cm, 0.0)

    def test_style_from_dict_ignores_junk(self):
        style = Style.from_dict({"font": "Arial", "size": "abc", "line_spacing": 0})
        self.assertEqual(style.font, "Arial")
        self.assertEqual(style.size, 12)
        self.assertEqual(style.line_spacing, 1.5)


class TestSplitWebApi(SplitTestCase):
    def setUp(self):
        super().setUp()
        from webapp.app import app

        app.config["TESTING"] = True
        self.app = app.test_client()

    def test_scan_endpoint(self):
        res = self.app.post("/api/split/scan", json={"path": str(self.epub)})
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.get_json()["total"], 3)

    def test_scan_requires_path(self):
        self.assertEqual(self.app.post("/api/split/scan", json={}).status_code, 400)

    def test_scan_missing_headings_asks_for_pattern(self):
        flat = self.tmp / "flat.txt"
        flat.write_text("Ни одного заголовка.\n", encoding="utf-8")
        res = self.app.post("/api/split/scan", json={"path": str(flat)})
        self.assertEqual(res.status_code, 422)
        body = res.get_json()
        self.assertTrue(body["need_pattern"])
        self.assertEqual(body["pattern"], corehead.DEFAULT_PATTERN)

    def test_browse_lists_book_files(self):
        data = self.app.get(
            "/api/browse", query_string={"path": str(self.tmp), "files": "epub,txt"}
        ).get_json()
        names = [f["name"] for f in data["files"]]
        self.assertIn("book.epub", names)
        self.assertIn("book.txt", names)

    def test_browse_hides_files_without_filter(self):
        data = self.app.get("/api/browse", query_string={"path": str(self.tmp)}).get_json()
        self.assertEqual(data["files"], [])

    def test_start_without_a_folder_name_writes_into_the_chosen_folder(self):
        """Сочинять папке имя незачем: человек уже выбрал, куда положить.

        Раньше пустое имя было отказом, и разбить файл «просто сюда» было
        нельзя — приходилось выдумывать название на каждый разбор.
        """
        res = self.app.post(
            "/api/split/start",
            json={"path": str(self.epub), "base": str(self.tmp), "folder": ""},
        )
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.get_json()["job"]["output_dir"],
                         str(self.tmp.resolve()))

    def test_start_still_requires_somewhere_to_put_them(self):
        """Папку не выбрали вовсе — это по-прежнему отказ, а не запись
        неизвестно куда."""
        res = self.app.post(
            "/api/split/start",
            json={"path": str(self.epub), "base": "", "folder": ""},
        )
        self.assertEqual(res.status_code, 400)

    def test_start_rejects_unknown_format(self):
        res = self.app.post(
            "/api/split/start",
            json={"path": str(self.epub), "base": str(self.tmp), "folder": "x", "format": "pdf"},
        )
        self.assertEqual(res.status_code, 400)

    def test_start_does_not_create_folder_on_bad_book(self):
        flat = self.tmp / "flat.txt"
        flat.write_text("Ни одного заголовка.\n", encoding="utf-8")
        res = self.app.post(
            "/api/split/start",
            json={"path": str(flat), "base": str(self.tmp), "folder": "should-not-exist"},
        )
        self.assertEqual(res.status_code, 422)
        self.assertFalse((self.tmp / "should-not-exist").exists())

    def test_full_split_job(self):
        from webapp.app import JOBS

        res = self.app.post(
            "/api/split/start",
            json={"path": str(self.epub), "base": str(self.tmp), "folder": "Книга",
                  "format": "txt"},
        )
        self.assertEqual(res.status_code, 200)
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=60)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertIsNone(job["error"])
        self.assertEqual(job["progress"]["stage"], "done")
        self.assertEqual(job["report"]["written"], 3)
        self.assertEqual(len(list((self.tmp / "Книга").glob("*.txt"))), 3)

    def test_full_split_job_docx(self):
        from webapp.app import JOBS

        res = self.app.post(
            "/api/split/start",
            json={"path": str(self.epub), "base": str(self.tmp), "folder": "Word",
                  "format": "docx"},
        )
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=60)
        self.assertEqual(len(list((self.tmp / "Word").glob("*.docx"))), 3)

    def test_job_snapshot_has_kind(self):
        from webapp.app import JOBS

        res = self.app.post(
            "/api/split/start",
            json={"path": str(self.epub), "base": str(self.tmp), "folder": "K"},
        )
        job_id = res.get_json()["job"]["id"]
        # Дожидаемся потока: иначе он пишет в папку, которую уже убрал tearDown.
        JOBS[job_id].thread.join(timeout=60)
        self.assertEqual(res.get_json()["job"]["kind"], "split")


class TestPerChapterSplitting(unittest.TestCase):
    """Деление главам поимённо и предпросмотр до записи.

    Раньше на всю книгу было одно поле «делить каждую главу на части».
    Забытая в нём двойка молча давала вдвое больше файлов, чем глав, и
    узнать об этом можно было только по готовой папке.
    """

    def setUp(self):
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)
        self.book = self.tmp / "книга.txt"
        lines = []
        for number in range(1, 6):
            lines.append(f"Глава {number}: Название {number}")
            lines += [f"Абзац {k} главы {number}. {LONG}" for k in range(1, 5)]
        self.book.write_text("\n\n".join(lines), encoding="utf-8")

    def look(self, **kw):
        from ops import split as split_op

        return split_op.look([str(self.book)], **kw)

    def test_a_book_without_settings_falls_apart_into_chapters(self):
        """Ради этого вкладка и существует."""
        data = self.look()
        self.assertEqual(data["found"], 5)
        self.assertEqual(data["total"], 5)
        self.assertEqual(len(data["names"]), 5)

    def test_one_chapter_can_be_cut_while_the_rest_stay_whole(self):
        data = self.look(pieces={"2": 3})
        self.assertEqual(data["found"], 5)
        self.assertEqual(data["total"], 7)
        self.assertEqual([row["parts"] for row in data["chapters"]],
                         [1, 3, 1, 1, 1])

    def test_the_named_chapter_wins_over_the_number_for_all(self):
        """Частное правило важнее общего — иначе поимённое деление
        нечем было бы отменить."""
        data = self.look(parts=2, pieces={"1": 1})
        self.assertEqual([row["parts"] for row in data["chapters"]],
                         [1, 2, 2, 2, 2])

    def test_the_table_says_how_big_each_chapter_is(self):
        """«Малая или большая» — то, ради чего в таблицу и смотрят."""
        for row in self.look()["chapters"]:
            with self.subTest(row["index"]):
                self.assertGreater(row["size"], 0)
                self.assertTrue(row["title"])

    def test_the_preview_shows_the_names_that_will_be_written(self):
        from ops import split as split_op

        out = self.tmp / "готово"
        want = self.look(pieces={"3": 2})["names"]
        split_op.run([str(self.book)], out, out_format=".txt",
                     pieces={"3": 2})
        got = sorted(path.stem for path in out.glob("*.txt"))
        self.assertEqual(sorted(want), got)

    def test_names_are_rebuilt_without_reading_the_book_again(self):
        """Предпросмотр перестраивается на каждую галочку, а книга в
        полторы тысячи глав читается с диска секунды."""
        from ops import split as split_op

        rows = [{"index": row["index"], "number": row["number"],
                 "title": row["title"]} for row in self.look()["chapters"]]
        # Файла на диске уже нет — значит, имена собраны по строкам.
        self.book.unlink()
        made = split_op.names(rows, pieces={"2": 2})
        self.assertEqual(len(made), 6)

    def test_the_part_number_is_written_the_chosen_way(self):
        from core import naming
        from ops import split as split_op

        rows = [{"index": 1, "number": 7, "title": "Глава 7: Имя"}]
        dot = split_op.names(rows, pieces={"1": 2},
                             fmt=naming.NameFormat.from_dict({}))
        word = split_op.names(
            rows, pieces={"1": 2},
            fmt=naming.NameFormat.from_dict({"part_style": "word"}))

        self.assertTrue(any("7.2" in name for name in dot), dot)
        self.assertTrue(any("Часть 2" in name for name in word), word)

    def test_the_ordinal_prefix_can_be_dropped(self):
        from core import naming
        from ops import split as split_op

        rows = [{"index": 1, "number": 7, "title": "Глава 7: Имя"}]
        made = split_op.names(rows, fmt=naming.NameFormat.from_dict({}),
                              seq=False)
        self.assertFalse(made[0].startswith("0001"), made)


class TestSplitOverHttp(unittest.TestCase):
    """Те же правила через маршруты: страница ходит только сюда."""

    def setUp(self):
        from webapp.app import app

        app.config["TESTING"] = True
        self.app = app.test_client()
        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)
        self.book = self.tmp / "книга.txt"
        lines = []
        for number in range(1, 4):
            lines.append(f"Глава {number}: Название {number}")
            lines += [f"Абзац {k} главы {number}. {LONG}" for k in range(1, 5)]
        self.book.write_text("\n\n".join(lines), encoding="utf-8")

    def test_scan_returns_the_chapters_and_the_names(self):
        data = self.app.post("/api/split/scan",
                             json={"path": str(self.book)}).get_json()
        self.assertEqual(len(data["chapters"]), 3)
        self.assertEqual(len(data["names"]), 3)

    def test_names_route_needs_no_file(self):
        data = self.app.post("/api/split/names", json={
            "chapters": [{"index": 1, "number": 1, "title": "Глава 1: Имя"}],
            "pieces": {"1": 2},
        }).get_json()
        self.assertEqual(len(data["names"]), 2)

    def test_the_name_format_does_not_collide_with_the_file_format(self):
        """`format` — расширение файла, имя собирается по `name_format`.
        Одно значение под двумя смыслами разошлось бы на первой правке."""
        data = self.app.post("/api/split/names", json={
            "chapters": [{"index": 1, "number": 5, "title": "Глава 5: Имя"}],
            "format": ".docx",
            "name_format": {"prefix": "Chapter", "title": False},
        }).get_json()
        self.assertEqual(data["names"], ["0001 - Chapter 5"])


if __name__ == "__main__":
    unittest.main(verbosity=2)


class TestVolumeCountsChaptersNotFiles(SplitTestCase):
    """«Проверить объём» на неразбитой книге.

    Общая проверка считает главой файл — она сделана для папки, где файл
    на главу и лежит. Здесь книга ещё одним файлом, и та же проверка
    честно отвечала «глав: 1, их слишком мало, чтобы говорить, какая
    выделяется», отвечая при этом не на тот вопрос.
    """

    def setUp(self):
        super().setUp()
        from webapp import app as web

        web.app.config["TESTING"] = True
        self.app = web.app.test_client()

    def test_it_sees_the_chapters_inside_the_book(self):
        book = make_txt(self.tmp / "целиком.txt", count=8)
        got = self.app.post("/api/split/volume",
                            json={"targets": [str(book)]}).get_json()
        self.assertGreaterEqual(got["chapters"], 8)

    def test_it_counts_exactly_what_the_split_will_make(self):
        """Проверка объёма и разбиение должны видеть одну и ту же книгу:
        разойдись они, и «выделяется третья» указывало бы не на ту."""
        book = make_txt(self.tmp / "целиком.txt", count=8)
        scan = self.app.post("/api/split/scan",
                             json={"targets": [str(book)]}).get_json()
        volume = self.app.post("/api/split/volume",
                               json={"targets": [str(book)]}).get_json()
        self.assertEqual(volume["chapters"], scan["found"])

    def test_the_old_check_still_sees_one_file(self):
        """Не поломка соседней проверки, а разные вопросы: там объём
        выбранных файлов, здесь — глав внутри книги."""
        book = make_txt(self.tmp / "целиком.txt", count=8)
        got = self.app.post("/api/stats",
                            json={"targets": [str(book)]}).get_json()
        self.assertEqual(got["chapters"], 1)

    def test_enough_chapters_means_it_can_actually_answer(self):
        """Ради этого ответа кнопку и жмут."""
        book = make_txt(self.tmp / "целиком.txt", count=8)
        got = self.app.post("/api/split/volume",
                            json={"targets": [str(book)]}).get_json()
        self.assertTrue(got["standout"]["enough"])

    def test_a_book_without_headings_asks_for_a_pattern(self):
        """Наугад не режем — и не считаем наугад тоже."""
        flat = self.tmp / "ровный.txt"
        flat.write_text("Ни одного заголовка.\n", encoding="utf-8")
        res = self.app.post("/api/split/volume", json={"targets": [str(flat)]})
        self.assertEqual(res.status_code, 422)
        self.assertTrue(res.get_json()["need_pattern"])


class TestTheTabOpensReadyToWork(unittest.TestCase):
    """Умолчания вкладки «Разбить».

    Разбивают книгу обычно одинаково: файл на главы, имена «Глава 99»,
    формат тот же, что у исходника. Всё это стояло иначе, и на каждый
    разбор приходилось снимать четыре галочки и выбирать формат заново.
    """

    @classmethod
    def setUpClass(cls):
        root = Path(__file__).resolve().parent.parent
        cls.html = (root / "webapp" / "static" / "index.html").read_text(
            encoding="utf-8")
        cls.js = (root / "webapp" / "static" / "tabs.js").read_text(
            encoding="utf-8")
        start = cls.html.index('id="tab-split"')
        cls.tab = cls.html[start:cls.html.index('id="tab-rename"')]

    def box(self, name):
        """Как в разметке записан этот флажок."""
        start = self.tab.index(f'id="{name}"')
        return self.tab[self.tab.rindex("<input", 0, start):
                        self.tab.index(">", start) + 1]

    def test_only_the_chapter_number_is_ticked(self):
        """«Глава 99» — и всё. Остальное человек поставит сам, если надо."""
        self.assertIn("checked", self.box("spNum"))
        for other in ("spSeq", "spPartNum", "spTitleOn"):
            self.assertNotIn("checked", self.box(other), other)

    def test_the_chapter_name_is_not_repeated_inside_the_text(self):
        """Название уже стоит в имени файла."""
        self.assertNotIn("checked", self.box("spHeadings"))

    def test_the_part_number_switches_itself_on_when_chapters_are_divided(self):
        """Без него все части главы получат одно имя, и запись разойдётся
        с предпросмотром приписками «(2)»."""
        start = self.js.index("function spApplyParts")
        body = self.js[start:self.js.index("\nfunction ", start + 1)]
        self.assertIn("spPartNum", body)

    def test_every_chapter_is_ticked_once_the_book_is_read(self):
        """Выбрали файл — значит, нужны все главы из него, а не часть."""
        start = self.js.index("async function spScan")
        body = self.js[start:self.js.index("window.spScan", start)]
        self.assertIn("spState.chosen.add(chapter.index)", body)

    def test_the_output_format_follows_the_source(self):
        """Разбивают вордовский файл — нужны вордовские главы."""
        self.assertIn("function spGuessFormat", self.js)
        start = self.js.index("async function spScan")
        self.assertIn("spGuessFormat",
                      self.js[start:self.js.index("window.spScan", start)])

    def test_the_preview_follows_the_chosen_format(self):
        """Предпросмотр обещал .txt, а на диск ложился .docx: он
        перерисовывался только при чтении с диска."""
        start = self.js.index("function spUpdateFinal")
        body = self.js[start:self.js.index("\n}", start)]
        self.assertIn("spDrawPreview", body)


class TestTheTabHidesWhatIsNotAsked(unittest.TestCase):
    """Свёрнутые карточки.

    Всё, что не основная работа вкладки, забирало глаз наравне с главным,
    и человек терялся на своей же вкладке.
    """

    @classmethod
    def setUpClass(cls):
        root = Path(__file__).resolve().parent.parent
        cls.html = (root / "webapp" / "static" / "index.html").read_text(
            encoding="utf-8")
        cls.js = (root / "webapp" / "static" / "tabs.js").read_text(
            encoding="utf-8")

    def card(self, name):
        start = self.html.index(f'id="{name}"')
        return self.html[self.html.rindex("<div", 0, start):
                         self.html.index(">", start) + 1]

    def test_the_side_cards_start_folded(self):
        for name in ("spVolCard", "spStyle", "spPrep", "spPreviewCard",
                     "spChaptersCard", "hdCard"):
            with self.subTest(name):
                self.assertIn("folded", self.card(name))
                self.assertIn("data-fold", self.card(name))

    def test_the_main_ones_are_not_folded(self):
        """Формат и место сохранения — это и есть работа вкладки."""
        for name in ("spOpts", "spPlace"):
            with self.subTest(name):
                self.assertNotIn("folded", self.card(name))

    def test_a_folded_card_only_shows_its_heading(self):
        self.assertIn(".card.folded > *:not(.foldhead){display:none}",
                      self.html.replace("\n", ""))

    def test_the_heading_opens_it(self):
        self.assertIn("foldhead", self.js)
        self.assertIn("classList.toggle('folded')", self.js)

    def test_divide_opens_what_divides(self):
        """Кнопки деления жили в чужой карточке выше, и найти их, не зная
        о них, было нельзя."""
        self.assertIn('id="spDivide"', self.html)
        self.assertIn("unfold('spChaptersCard')", self.js)
