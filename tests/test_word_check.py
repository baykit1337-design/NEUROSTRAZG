"""Тесты вкладок «В Word» и «Проверка текста», диалога и таймаута.

Разделы 1, 2, 4 и 5 ТЗ.
"""

from __future__ import annotations

import sys
import threading
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from mvl import client as client_mod  # noqa: E402
from mvl import nativedialog, textcheck, toword  # noqa: E402
from mvl.word import Style  # noqa: E402

LONG = "Длинный абзац русского текста для проверки объёма. " * 12


def make_chapters(root: Path) -> Path:
    folder = root / "главы"
    folder.mkdir()
    (folder / "0001 - Глава 1. Начало.txt").write_text(
        "# Внутренний заголовок\n\n"
        "Абзац с **жирным**, *курсивом* и `кодом`.\n\n"
        "> Цитата\n\n---\n\n*\n\nПосле разделителя [ссылка](http://x.ru) внутри.\n",
        encoding="utf-8",
    )
    (folder / "0002 - Глава 2. Дальше.txt").write_text(
        f"{LONG}\n\nВторой абзац.\n", encoding="utf-8"
    )
    return folder


class WordTestCase(unittest.TestCase):
    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)
        self.folder = make_chapters(self.tmp)


class TestConvert(WordTestCase):
    def test_single_document(self):
        out = self.tmp / "книга.docx"
        report = toword.convert(self.folder, out, mode=toword.MODE_SINGLE)
        self.assertEqual((report.written, report.failed), (2, 0))
        self.assertTrue(out.is_file())

    def test_per_chapter_on_plain_files_says_there_is_nothing_to_do(self):
        """Молча копировать готовые файлы нельзя — надо сказать об этом."""
        out = self.tmp / "по главам"
        with self.assertRaises(toword.NothingToDo) as ctx:
            toword.convert(self.folder, out, mode=toword.MODE_PER_CHAPTER)
        self.assertIn("уже отдельные файлы", str(ctx.exception))

    def test_headings_use_heading_1(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert(self.folder, out)
        styles = [p.style.name for p in Document(str(out)).paragraphs if p.text.strip()]
        self.assertEqual(styles[0], "Heading 1")

    def test_sequence_number_is_not_in_the_heading(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert(self.folder, out)
        first = next(p.text for p in Document(str(out)).paragraphs if p.text.strip())
        self.assertEqual(first, "Глава 1. Начало")

    def test_no_table_of_contents_and_no_title_page(self):
        """Оглавление и титул убраны из ТЗ — их не должно быть и в настройках."""
        self.assertFalse(hasattr(Style(), "table_of_contents"))
        self.assertFalse(hasattr(Style(), "title_page"))

    def test_style_applied(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert(self.folder, out, style=Style(font="Arial", size=14))
        document = Document(str(out))
        self.assertEqual(document.styles["Normal"].font.name, "Arial")

    def test_markdown_is_rendered_not_carried_as_symbols(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert(self.folder, out)
        document = Document(str(out))
        text = "\n".join(p.text for p in document.paragraphs)

        self.assertNotIn("**", text)
        self.assertNotIn("`", text)
        self.assertNotIn("](", text)

        runs = [r for p in document.paragraphs for r in p.runs]
        self.assertTrue(any(r.bold and r.text == "жирным" for r in runs))
        self.assertTrue(any(r.italic and r.text == "курсивом" for r in runs))

    def test_inner_heading_becomes_a_heading(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert(self.folder, out)
        headings = [p.text for p in Document(str(out)).paragraphs
                    if p.style.name.startswith("Heading")]
        self.assertIn("Внутренний заголовок", headings)

    def test_scene_separator_becomes_the_chosen_form(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert(self.folder, out)
        texts = [p.text.strip() for p in Document(str(out)).paragraphs]
        self.assertIn("* * *", texts)

    def test_broken_file_is_reported_with_file_step_and_error(self):
        """Молчаливый отказ недопустим."""
        (self.folder / "0003 - Глава 3. Битая.docx").write_text("не docx", encoding="utf-8")
        report = toword.convert(self.folder, self.tmp / "к.docx")

        self.assertEqual(report.failed, 1)
        failure = report.failures[0]
        self.assertEqual(failure.file, "0003 - Глава 3. Битая.docx")
        self.assertTrue(failure.step)
        self.assertIn("Error", failure.error)

    def test_one_broken_file_does_not_stop_the_rest(self):
        (self.folder / "0003 - Глава 3. Битая.docx").write_text("не docx", encoding="utf-8")
        report = toword.convert(self.folder, self.tmp / "к.docx")
        self.assertEqual(report.written, 2)

    def test_all_files_broken_raises_with_reason(self):
        empty = self.tmp / "битые"
        empty.mkdir()
        (empty / "a.docx").write_text("не docx", encoding="utf-8")
        with self.assertRaises(toword.ConvertError) as ctx:
            toword.convert(empty, self.tmp / "x.docx")
        self.assertIn("a.docx", str(ctx.exception))

    def test_missing_folder_reports_clearly(self):
        with self.assertRaises(toword.ConvertError):
            toword.convert(self.tmp / "нет такой", self.tmp / "x.docx")

    def test_cancel(self):
        cancel = threading.Event()
        cancel.set()
        from mvl.booksplit import Cancelled

        with self.assertRaises(Cancelled):
            toword.convert(self.folder, self.tmp / "x.docx", cancel=cancel)

    def test_progress(self):
        seen = []
        toword.convert(self.folder, self.tmp / "x.docx",
                       on_progress=lambda d, t: seen.append((d, t)))
        self.assertEqual(seen, [(1, 2), (2, 2)])


class TestTextCheck(unittest.TestCase):
    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)

    def write(self, name: str, text: str) -> Path:
        path = self.tmp / name
        path.write_text(text, encoding="utf-8")
        return path

    def kinds_found(self, report) -> set:
        return {f.kind for f in report.findings}

    def test_cjk(self):
        self.write("a.txt", f"Русский текст 修炼 внутри.\n{LONG}")
        self.assertIn("cjk", self.kinds_found(textcheck.check(self.tmp, ["cjk"])))

    def test_korean_and_japanese(self):
        self.write("a.txt", f"Текст 한글 и ひらがな и カタカナ.\n{LONG}")
        self.assertIn("cjk", self.kinds_found(textcheck.check(self.tmp, ["cjk"])))

    def test_markdown_leftovers(self):
        self.write("a.txt", f"Тут **жирный** остался.\n{LONG}")
        self.assertIn("markdown", self.kinds_found(textcheck.check(self.tmp, ["markdown"])))

    def test_latin_words_are_counted(self):
        self.write("a.txt", f"Слово manager осталось. И снова manager тут.\n{LONG}")
        report = textcheck.check(self.tmp, ["latin"])
        self.assertEqual(report.latin_words[0], {"word": "manager", "count": 2})

    def test_latin_ignores_urls_and_entities(self):
        self.write("a.txt", f"Русский текст http://example.com и &nbsp; тут.\n{LONG}")
        report = textcheck.check(self.tmp, ["latin"])
        self.assertEqual(report.latin_words, [])

    def test_latin_only_inside_russian_text(self):
        self.write("a.txt", "Pure english line without any cyrillic\n" + LONG)
        report = textcheck.check(self.tmp, ["latin"])
        self.assertNotIn("english", [row["word"] for row in report.latin_words])

    def test_whitelist_excludes_words(self):
        self.write("a.txt", f"Слово manager осталось.\n{LONG}")
        self.write("whitelist.txt", "manager\n")
        report = textcheck.check(self.tmp, ["latin"])
        self.assertEqual(report.latin_words, [])

    def test_whitelist_file_is_not_checked_as_a_chapter(self):
        self.write("a.txt", LONG)
        self.write("whitelist.txt", "manager\n")
        self.assertEqual(textcheck.check(self.tmp, ["size"]).files_checked, 1)

    def test_model_traces(self):
        for text in ["Note: something", "Вот перевод главы", "As an AI I cannot",
                     "I hope this helps", "Translator's note here"]:
            with self.subTest(text=text):
                self.write("a.txt", f"{text}\n{LONG}")
                self.assertIn("model", self.kinds_found(textcheck.check(self.tmp, ["model"])))

    def test_broken_encoding_and_html(self):
        for text in ["Битый � символ", "Тут &nbsp; сущность", "Тут <p>тег</p>"]:
            with self.subTest(text=text):
                self.write("a.txt", f"{text}\n{LONG}")
                self.assertIn("broken", self.kinds_found(textcheck.check(self.tmp, ["broken"])))

    def test_loop_detected_at_three_repeats(self):
        self.write("a.txt", "Повтор.\nПовтор.\nПовтор.\n" + LONG)
        report = textcheck.check(self.tmp, ["loop"])
        self.assertIn("loop", self.kinds_found(report))

    def test_two_repeats_is_not_a_loop(self):
        self.write("a.txt", "Повтор.\nПовтор.\n" + LONG)
        self.assertEqual(textcheck.check(self.tmp, ["loop"]).findings, [])

    def test_short_file(self):
        self.write("a.txt", "Кратко.")
        self.assertIn("size", self.kinds_found(textcheck.check(self.tmp, ["size"])))

    def test_file_far_from_the_median(self):
        for index in range(5):
            self.write(f"n{index}.txt", LONG)
        self.write("big.txt", LONG * 20)
        report = textcheck.check(self.tmp, ["size"])
        self.assertIn("big.txt", {f.file for f in report.findings})

    def test_unbalanced_quotes_and_brackets(self):
        self.write("a.txt", f"Открыл «кавычку и не закрыл.\n{LONG}")
        self.assertIn("pairs", self.kinds_found(textcheck.check(self.tmp, ["pairs"])))

    def test_balanced_text_is_clean(self):
        self.write("a.txt", f"Всё «закрыто» и (тоже).\n{LONG}")
        self.assertEqual(textcheck.check(self.tmp, ["pairs"]).findings, [])

    def test_findings_carry_file_line_and_fragment(self):
        self.write("a.txt", f"{LONG}\nЗдесь 修炼 иероглиф.\n")
        finding = textcheck.check(self.tmp, ["cjk"]).findings[0]
        self.assertEqual(finding.file, "a.txt")
        self.assertEqual(finding.line, 2)
        self.assertIn("修炼", finding.fragment)

    def test_summary_counts_by_kind(self):
        self.write("a.txt", f"Тут 修炼 и **жирный**.\n{LONG}")
        report = textcheck.check(self.tmp, ["cjk", "markdown"])
        self.assertEqual(report.summary, {"cjk": 1, "markdown": 1})

    def test_files_without_findings_are_not_listed(self):
        self.write("clean.txt", LONG)
        self.write("dirty.txt", f"Тут 修炼.\n{LONG}")
        report = textcheck.check(self.tmp, ["cjk"])
        self.assertEqual({f.file for f in report.findings}, {"dirty.txt"})
        self.assertEqual(report.files_with_findings, 1)

    def test_single_file_target(self):
        path = self.write("a.txt", f"Тут 修炼.\n{LONG}")
        self.assertEqual(textcheck.check(path, ["cjk"]).files_checked, 1)

    def test_unknown_kind_rejected(self):
        self.write("a.txt", LONG)
        with self.assertRaises(textcheck.CheckError):
            textcheck.check(self.tmp, ["нет такой"])

    def test_missing_target(self):
        with self.assertRaises(textcheck.CheckError):
            textcheck.check(self.tmp / "нет")

    def test_unreadable_file_does_not_stop_the_rest(self):
        self.write("good.txt", f"Тут 修炼.\n{LONG}")
        (self.tmp / "bad.docx").write_text("не docx", encoding="utf-8")
        report = textcheck.check(self.tmp, ["cjk"])
        self.assertEqual(len(report.unreadable), 1)
        self.assertIn("cjk", self.kinds_found(report))

    def test_cancel(self):
        from mvl.booksplit import Cancelled

        self.write("a.txt", LONG)
        cancel = threading.Event()
        cancel.set()
        with self.assertRaises(Cancelled):
            textcheck.check(self.tmp, cancel=cancel)

    def test_report_text_lists_everything(self):
        self.write("a.txt", f"Тут 修炼 и слово manager.\n{LONG}")
        text = textcheck.report_text(textcheck.check(self.tmp))
        self.assertIn("Проверка текста", text)
        self.assertIn("Сводка по типам", text)
        self.assertIn("manager", text)


class TestTimeout(unittest.TestCase):
    """Раздел 1: страница ~220 КБ не успевала прийти за 30 секунд."""

    def test_defaults_from_spec(self):
        self.assertEqual(client_mod.TIMEOUT, 120)
        self.assertEqual(client_mod.CONNECT_TIMEOUT, 15)

    def test_client_keeps_both_timeouts(self):
        client = client_mod.Client(timeout=90, connect_timeout=7)
        self.assertEqual((client.timeout, client.connect_timeout), (90, 7))

    def test_incomplete_response_named_precisely(self):
        class Response:
            headers = {"Content-Length": "220028"}
            content = b"x" * 22000

        message = client_mod._incomplete(Response())
        self.assertIn("22000", message)
        self.assertIn("неполный", message)

    def test_complete_response_is_not_flagged(self):
        class Response:
            headers = {"Content-Length": "100"}
            content = b"x" * 100

        self.assertIsNone(client_mod._incomplete(Response()))

    def test_compressed_response_is_not_flagged(self):
        """Сжатый ответ распакован — длина законно расходится с заголовком."""

        class Response:
            headers = {"Content-Length": "5000", "Content-Encoding": "gzip"}
            content = b"x" * 20000

        self.assertIsNone(client_mod._incomplete(Response()))

    def test_torn_connection_is_not_called_a_timeout(self):
        message = client_mod._describe(Exception("Transfer closed with outstanding read data"))
        self.assertIn("неполный", message)

    def test_plain_timeout_keeps_its_own_wording(self):
        self.assertNotIn("неполный", client_mod._describe(Exception("Operation timed out")))


class TestNativeDialog(unittest.TestCase):
    """Раздел 2. Без графической оболочки диалог обязан отказать внятно."""

    def test_available_returns_bool(self):
        self.assertIsInstance(nativedialog.available(), bool)

    def test_no_display_raises_dialog_unavailable(self):
        if nativedialog.available():
            self.skipTest("Графическая оболочка есть — отказ не воспроизвести")
        with self.assertRaises(nativedialog.DialogUnavailable):
            nativedialog.ask_directory()


class TestWordCheckWebApi(WordTestCase):
    def setUp(self):
        super().setUp()
        from webapp.app import app

        app.config["TESTING"] = True
        self.app = app.test_client()

    def test_word_scan(self):
        res = self.app.post("/api/word/scan", json={"folder_in": str(self.folder)})
        self.assertEqual(res.status_code, 200)
        self.assertEqual(res.get_json()["total"], 2)

    def test_word_rejects_unknown_mode(self):
        res = self.app.post("/api/word/start", json={
            "folder_in": str(self.folder), "base": str(self.tmp),
            "name": "К", "mode": "нет такого"})
        self.assertEqual(res.status_code, 400)

    def test_word_single_job(self):
        from webapp.app import JOBS

        res = self.app.post("/api/word/start", json={
            "folder_in": str(self.folder), "base": str(self.tmp),
            "name": "Книга", "mode": "single"})
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=60)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertIsNone(job["error"])
        self.assertEqual(job["report"]["written"], 2)
        self.assertTrue((self.tmp / "Книга.docx").is_file())

    def test_word_per_chapter_on_plain_files_reports_nothing_to_do(self):
        from webapp.app import JOBS

        res = self.app.post("/api/word/start", json={
            "targets": [str(self.folder)], "base": str(self.tmp),
            "name": "Главы", "mode": "per_chapter"})
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=60)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertIn("раскладывать нечего", job["error"])

    def test_check_job_and_report_download(self):
        from webapp.app import JOBS

        (self.folder / "плохая.txt").write_text(f"Тут 修炼.\n{LONG}", encoding="utf-8")
        res = self.app.post("/api/check/start", json={"target": str(self.folder)})
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=60)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertIsNone(job["error"])
        self.assertGreater(job["report"]["total"], 0)

        report = self.app.get(f"/api/check/{job_id}/report")
        self.assertEqual(report.status_code, 200)
        self.assertIn("Проверка текста", report.get_data(as_text=True))

    def test_check_requires_target(self):
        self.assertEqual(self.app.post("/api/check/start", json={}).status_code, 400)

    def test_check_requires_at_least_one_kind(self):
        res = self.app.post("/api/check/start",
                            json={"target": str(self.folder), "kinds": []})
        self.assertEqual(res.status_code, 400)

    def test_pick_available_endpoint(self):
        res = self.app.get("/api/pick/available")
        self.assertEqual(res.status_code, 200)
        self.assertIn("available", res.get_json())

    def test_pick_unknown_kind(self):
        self.assertEqual(self.app.post("/api/pick/чтото", json={}).status_code, 404)

    def test_pick_without_display_reports_fallback(self):
        if nativedialog.available():
            self.skipTest("Графическая оболочка есть")
        res = self.app.post("/api/pick/folder", json={})
        self.assertEqual(res.status_code, 503)
        self.assertTrue(res.get_json()["fallback"])

    def test_download_rejects_tiny_timeout(self):
        res = self.app.post("/api/start", json={
            "novel": {"code": 1, "name": "x", "total_chapters": 1},
            "base": str(self.tmp), "folder": "x", "timeout": 1})
        self.assertEqual(res.status_code, 400)


if __name__ == "__main__":
    unittest.main(verbosity=2)
