"""Тесты правок из ТЗ v4.

Главное здесь — epub больше не читается как бинарник: из этой ошибки росли
сразу две проблемы, и обе проверяются ниже.
"""

from __future__ import annotations

import sys
import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from mvl import cleanup, source, textcheck, textprep, toword  # noqa: E402
from mvl.textprep import PrepOptions  # noqa: E402
from tests.test_split import make_epub  # noqa: E402

LONG = "Обычный абзац русского текста, достаточно длинный для проверки. " * 6


class V4TestCase(unittest.TestCase):
    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)
        self.epub = make_epub(self.tmp / "book.epub", count=3)


class TestEpubReading(V4TestCase):
    """КРИТИЧНО из ТЗ: epub это ZIP, а не текст."""

    def test_epub_gives_chapters_not_raw_bytes(self):
        chapters = source.load_chapters(self.epub)
        self.assertEqual(len(chapters), 3)
        self.assertTrue(all(c.paragraphs for c in chapters))

    def test_epub_text_has_no_replacement_characters(self):
        """Сырые байты ZIP давали мусор вида `??v????A` — его быть не должно."""
        text = "\n".join(p for c in source.load_chapters(self.epub) for p in c.paragraphs)
        self.assertNotIn("�", text)

    def test_check_on_epub_finds_no_broken_encoding(self):
        report = textcheck.check(self.epub, ["broken"])
        self.assertEqual(report.findings, [])

    def test_check_accepts_epub_as_a_single_file(self):
        report = textcheck.check(self.epub, ["cjk"])
        self.assertEqual(report.files_checked, 1)

    def test_word_accepts_epub(self):
        """Раньше вкладка отвечала «в папке нет файлов .txt, .md или .docx»."""
        info = toword.scan([str(self.epub)])
        self.assertEqual(info["total"], 3)
        self.assertTrue(info["multi_chapter"])

    def test_epub_to_single_docx(self):
        out = self.tmp / "книга.docx"
        report = toword.convert([str(self.epub)], out, mode=toword.MODE_SINGLE)
        self.assertEqual(report.written, 3)
        self.assertTrue(out.is_file())

    def test_epub_split_into_separate_docx(self):
        out = self.tmp / "главы"
        report = toword.convert([str(self.epub)], out, mode=toword.MODE_PER_CHAPTER)
        self.assertEqual(report.written, 3)
        self.assertEqual(len(list(out.glob("*.docx"))), 3)

    def test_unknown_extension_is_refused_not_read_as_text(self):
        bad = self.tmp / "book.pdf"
        bad.write_bytes(b"%PDF-1.4 \x00\x01\x02")
        with self.assertRaises(source.SourceError):
            source.collect_sources([str(bad)])

    def test_check_refuses_unknown_single_file(self):
        bad = self.tmp / "book.pdf"
        bad.write_bytes(b"%PDF-1.4 \x00\x01\x02")
        with self.assertRaises(textcheck.CheckError):
            textcheck.check(bad)


class TestMultipleTargets(V4TestCase):
    def test_files_and_folders_mixed(self):
        folder = self.tmp / "тексты"
        folder.mkdir()
        (folder / "a.txt").write_text(LONG, encoding="utf-8")
        files = source.collect_sources([str(self.epub), str(folder)])
        self.assertEqual(len(files), 2)

    def test_duplicates_are_not_counted_twice(self):
        files = source.collect_sources([str(self.epub), str(self.epub)])
        self.assertEqual(len(files), 1)

    def test_check_takes_several_targets(self):
        second = make_epub(self.tmp / "second.epub", count=2)
        report = textcheck.check([str(self.epub), str(second)], ["cjk"])
        self.assertEqual(report.files_checked, 2)


class TestTextPrep(unittest.TestCase):
    """Раздел 6: качество docx на выходе."""

    def test_duplicate_title_removed(self):
        title = "Глава 209. Частичное приручение паука? (1)"
        blocks = textprep.prepare([title, title, "Текст главы."], title)
        self.assertEqual([b.text for b in blocks], ["Текст главы."])

    def test_title_matched_without_chapter_prefix(self):
        blocks = textprep.prepare(
            ["Глава 209. Название", "Текст."], "Название"
        )
        self.assertEqual([b.text for b in blocks], ["Текст."])

    def test_title_matched_ignoring_case_and_spaces(self):
        blocks = textprep.prepare(["  НАЗВАНИЕ  ", "Текст."], "Название")
        self.assertEqual([b.text for b in blocks], ["Текст."])

    def test_only_leading_paragraphs_are_checked(self):
        """Название дальше по тексту — часть повествования, не трогаем."""
        paragraphs = ["А", "Б", "В", "Название", "Г"]
        blocks = textprep.prepare(paragraphs, "Название")
        self.assertIn("Название", [b.text for b in blocks])

    def test_title_removal_can_be_switched_off(self):
        blocks = textprep.prepare(["Название", "Текст."], "Название",
                                  PrepOptions(strip_title=False))
        self.assertEqual(len(blocks), 2)

    def test_consecutive_scene_breaks_collapse(self):
        blocks = textprep.prepare(["А", "*", "*", "*", "Б"], "")
        self.assertEqual([b.text for b in blocks], ["А", "* * *", "Б"])

    def test_scene_break_not_first_or_last(self):
        blocks = textprep.prepare(["*", "А", "*"], "")
        self.assertEqual([b.text for b in blocks], ["А"])

    def test_scene_break_forms(self):
        cases = {
            textprep.SCENE_STARS: "* * *",
            textprep.SCENE_DASHES: "— — —",
            textprep.SCENE_BLANK: "",
        }
        for style, expected in cases.items():
            blocks = textprep.prepare(["А", "*", "Б"], "", PrepOptions(scene_style=style))
            self.assertEqual(blocks[1].text, expected, style)

    def test_scene_break_kept_as_is(self):
        blocks = textprep.prepare(["А", "※", "Б"], "", PrepOptions(scene_style="keep"))
        self.assertEqual(blocks[1].text, "※")

    def test_various_separator_characters(self):
        for mark in ["*", "＊", "※", "···", "---", "—", "***", "* * *"]:
            with self.subTest(mark=mark):
                self.assertTrue(textprep.is_scene_break(mark))

    def test_ordinary_text_is_not_a_separator(self):
        self.assertFalse(textprep.is_scene_break("Обычный текст"))

    def test_empty_paragraphs_dropped(self):
        blocks = textprep.prepare(["А", "", "   ", "Б"], "")
        self.assertEqual([b.text for b in blocks], ["А", "Б"])

    def test_system_messages_are_marked_not_removed(self):
        blocks = textprep.prepare(["{Обнаружен Потенциал}", "[Получено достижение]"], "")
        self.assertEqual(len(blocks), 2)
        self.assertTrue(all(b.kind == textprep.KIND_SYSTEM for b in blocks))

    def test_defaults_from_spec(self):
        options = PrepOptions()
        self.assertEqual(options.align, "left")
        self.assertEqual(options.first_line_indent_cm, 0.0)
        self.assertTrue(options.strip_title)
        self.assertFalse(options.italic_system)


class TestDocxOutput(V4TestCase):
    def test_everything_is_left_aligned_except_separators(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        out = self.tmp / "к.docx"
        toword.convert([str(self.epub)], out)
        for paragraph in Document(str(out)).paragraphs:
            if not paragraph.text.strip() or paragraph.style.name.startswith("Heading"):
                continue
            if paragraph.text.strip() in ("* * *", "— — —"):
                self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.CENTER)
            else:
                self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT, paragraph.text[:40])

    def test_first_line_indent_is_zero_by_default(self):
        from docx import Document

        out = self.tmp / "к.docx"
        toword.convert([str(self.epub)], out)
        for paragraph in Document(str(out)).paragraphs:
            indent = paragraph.paragraph_format.first_line_indent
            if indent is not None:
                self.assertEqual(indent, 0, paragraph.text[:40])

    def test_system_messages_align_with_body_text(self):
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        folder = self.tmp / "сист"
        folder.mkdir()
        (folder / "Глава 1.txt").write_text(
            f"{LONG}\n\n{{Обнаружен Потенциал эволюции}}\n\n[Получено достижение]\n",
            encoding="utf-8",
        )
        out = self.tmp / "s.docx"
        toword.convert([str(folder)], out)
        for paragraph in Document(str(out)).paragraphs:
            if paragraph.text.strip().startswith(("{", "[")):
                self.assertEqual(paragraph.alignment, WD_ALIGN_PARAGRAPH.LEFT)

    def test_system_messages_can_be_italic(self):
        from docx import Document

        folder = self.tmp / "сист"
        folder.mkdir()
        (folder / "Глава 1.txt").write_text(
            f"{LONG}\n\n{{Обнаружен Потенциал}}\n", encoding="utf-8"
        )
        out = self.tmp / "s.docx"
        toword.convert([str(folder)], out, prep=PrepOptions(italic_system=True))
        runs = [r for p in Document(str(out)).paragraphs for r in p.runs
                if p.text.strip().startswith("{")]
        self.assertTrue(runs and all(r.italic for r in runs))


class TestCleanup(unittest.TestCase):
    """Раздел 3.3: кнопка «Очистить»."""

    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)
        self.src = self.tmp / "in"
        self.src.mkdir()

    def write(self, text: str) -> Path:
        path = self.src / "a.txt"
        path.write_text(text, encoding="utf-8")
        return path

    def clean(self, text: str, kinds) -> str:
        self.write(text)
        out = self.tmp / "out"
        cleanup.clean(self.src, kinds, out)
        return (out / "a.txt").read_text(encoding="utf-8")

    def test_markdown_removed_but_text_kept(self):
        result = self.clean("Текст с **жирным** и `кодом`.", ["markdown"])
        self.assertIn("жирным", result)
        self.assertNotIn("**", result)
        self.assertNotIn("`", result)

    def test_markdown_heading_and_quote(self):
        result = self.clean("## Заголовок\n> Цитата", ["markdown"])
        self.assertIn("Заголовок", result)
        self.assertNotIn("##", result)
        self.assertNotIn(">", result)

    def test_html_tags_and_entities(self):
        result = self.clean("<p>Текст</p> и &amp; и &nbsp; тут", ["html"])
        self.assertNotIn("<p>", result)
        self.assertIn("&", result)
        self.assertNotIn("&amp;", result)
        self.assertNotIn("&nbsp;", result)

    def test_nbsp_becomes_a_space(self):
        result = self.clean("А Б", ["nbsp"])
        self.assertNotIn(" ", result)
        self.assertIn("А Б", result)

    def test_broken_characters_removed(self):
        self.assertNotIn("�", self.clean("Текст � тут", ["broken"]))

    def test_model_traces_removed_by_line(self):
        result = self.clean("Обычный текст.\nNote: model trace.\nЕщё текст.", ["model"])
        self.assertNotIn("Note:", result)
        self.assertIn("Обычный текст.", result)
        self.assertIn("Ещё текст.", result)

    def test_repeated_paragraphs_collapse_to_one(self):
        result = self.clean("Повтор.\nПовтор.\nПовтор.\nКонец.", ["dupes"])
        self.assertEqual(result.count("Повтор."), 1)

    def test_extra_blank_lines_collapse(self):
        result = self.clean("А\n\n\n\n\nБ", ["blanks"])
        self.assertNotIn("\n\n\n", result)

    def test_cjk_and_latin_are_never_touched(self):
        """Их надо смотреть глазами — машинально удалять нельзя."""
        result = self.clean("Текст 修炼 и manager тут.", list(cleanup.ALL_KINDS))
        self.assertIn("修炼", result)
        self.assertIn("manager", result)

    def test_preview_counts_without_writing(self):
        self.write("**жирный** и <p>тег</p>")
        before = set(self.tmp.iterdir())
        result = cleanup.preview(self.src, ["markdown", "html"])
        self.assertGreater(result["total"], 0)
        self.assertEqual(set(self.tmp.iterdir()), before)

    def test_originals_are_not_touched(self):
        original = "**жирный** остаётся в оригинале"
        self.write(original)
        cleanup.clean(self.src, ["markdown"], self.tmp / "out")
        self.assertEqual((self.src / "a.txt").read_text(encoding="utf-8"), original)

    def test_report_says_what_was_fixed(self):
        self.write("**жирный** и <p>тег</p> и  ")
        report = cleanup.clean(self.src, ["markdown", "html", "nbsp"], self.tmp / "out")
        counts = {row["kind"]: row["count"] for row in report.as_dict()["counts"]}
        self.assertGreater(counts.get("markdown", 0), 0)
        self.assertGreater(counts.get("html", 0), 0)
        self.assertEqual(counts.get("nbsp"), 1)

    def test_empty_selection_refused(self):
        self.write("текст")
        with self.assertRaises(cleanup.CleanError):
            cleanup.clean(self.src, [], self.tmp / "out")

    def test_unknown_kind_refused(self):
        self.write("текст")
        with self.assertRaises(cleanup.CleanError):
            cleanup.clean(self.src, ["нет такой"], self.tmp / "out")


class TestV4WebApi(V4TestCase):
    def setUp(self):
        super().setUp()
        from webapp.app import app

        app.config["TESTING"] = True
        self.app = app.test_client()

    def test_word_scan_accepts_epub(self):
        res = self.app.post("/api/word/scan", json={"targets": [str(self.epub)]})
        self.assertEqual(res.status_code, 200)
        body = res.get_json()
        self.assertEqual(body["total"], 3)
        self.assertTrue(body["multi_chapter"])

    def test_word_scan_accepts_several_files(self):
        second = make_epub(self.tmp / "second.epub", count=2)
        res = self.app.post("/api/word/scan",
                            json={"targets": [str(self.epub), str(second)]})
        self.assertEqual(res.get_json()["total"], 5)

    def test_epub_to_docx_job(self):
        from webapp.app import JOBS

        res = self.app.post("/api/word/start", json={
            "targets": [str(self.epub)], "base": str(self.tmp),
            "name": "Книга", "mode": "single"})
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=120)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertIsNone(job["error"])
        self.assertEqual(job["report"]["written"], 3)
        self.assertTrue((self.tmp / "Книга.docx").is_file())

    def test_check_on_epub_job(self):
        from webapp.app import JOBS

        res = self.app.post("/api/check/start",
                            json={"targets": [str(self.epub)], "kinds": ["broken"]})
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=120)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertEqual(job["report"]["total"], 0)

    def test_clean_preview_endpoint(self):
        folder = self.tmp / "гряз"
        folder.mkdir()
        (folder / "a.txt").write_text("**жирный** и <p>тег</p>", encoding="utf-8")
        res = self.app.post("/api/clean/preview",
                            json={"targets": [str(folder)], "kinds": ["markdown", "html"]})
        self.assertEqual(res.status_code, 200)
        self.assertGreater(res.get_json()["total"], 0)

    def test_clean_job(self):
        from webapp.app import JOBS

        folder = self.tmp / "гряз"
        folder.mkdir()
        (folder / "a.txt").write_text("**жирный** и <p>тег</p>", encoding="utf-8")

        res = self.app.post("/api/clean/start", json={
            "targets": [str(folder)], "base": str(self.tmp), "folder": "Чисто",
            "kinds": ["markdown", "html"]})
        job_id = res.get_json()["job"]["id"]
        JOBS[job_id].thread.join(timeout=60)

        job = self.app.get(f"/api/job/{job_id}").get_json()["job"]
        self.assertIsNone(job["error"])
        self.assertGreater(job["report"]["total"], 0)
        self.assertTrue((self.tmp / "Чисто" / "a.txt").is_file())

    def test_clean_requires_kinds(self):
        res = self.app.post("/api/clean/start", json={
            "targets": [str(self.tmp)], "base": str(self.tmp),
            "folder": "x", "kinds": []})
        self.assertEqual(res.status_code, 400)

    def test_timeout_up_to_300_is_accepted(self):
        res = self.app.post("/api/start", json={
            "novel": {"code": 1, "name": "x", "total_chapters": 1},
            "base": str(self.tmp), "folder": "т", "timeout": 300})
        # Дальше запрос уходит в сеть, но проверку таймаута он уже прошёл.
        self.assertNotEqual(res.status_code, 400)

    def test_timeout_above_300_refused(self):
        res = self.app.post("/api/start", json={
            "novel": {"code": 1, "name": "x", "total_chapters": 1},
            "base": str(self.tmp), "folder": "т", "timeout": 301})
        self.assertEqual(res.status_code, 400)

    def test_pick_any_kind_exists(self):
        from mvl import nativedialog

        if nativedialog.available():
            self.skipTest("Графическая оболочка есть")
        res = self.app.post("/api/pick/any", json={})
        self.assertEqual(res.status_code, 503)


if __name__ == "__main__":
    unittest.main(verbosity=2)
