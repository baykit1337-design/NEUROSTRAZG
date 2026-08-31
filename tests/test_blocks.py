"""Деление сплошного .docx по разметке, а не по тексту.

Документы здесь собираются на лету и повторяют то, что делает Word, когда
в него вставляют главу со страницы сайта, не снимая форматирования.
"""

from __future__ import annotations

import sys
import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from core import blocks  # noqa: E402
from core.blocks import NoBlocksFound  # noqa: E402

SIDES = ("left", "right")


def _border(paragraph, sides):
    """Ставит абзацу рамку с указанными сторонами; остальные — `nil`.

    Word пишет все четыре стороны всегда: у той, которой нет, стоит `nil`.
    Повторяем это, иначе проверка шла бы по документу, какого не бывает.
    """
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    frame = OxmlElement("w:pBdr")
    for side in ("top", "left", "bottom", "right"):
        edge = OxmlElement(f"w:{side}")
        if side in sides:
            edge.set(qn("w:val"), "single")
            edge.set(qn("w:sz"), "2")
            edge.set(qn("w:color"), "898073")
        else:
            edge.set(qn("w:val"), "nil")
        frame.append(edge)
    paragraph._p.get_or_add_pPr().append(frame)


def make_docx(path: Path, rows) -> Path:
    """Документ по описанию: (текст, стороны рамки) на абзац."""
    from docx import Document

    document = Document()
    for text, sides in rows:
        paragraph = document.add_paragraph(text)
        if sides:
            _border(paragraph, sides)
    document.save(str(path))
    return path


def chapter_box(number: int, size: int = 3):
    """Глава, обведённая рамкой поабзацно, — как её кладёт Word."""
    return [(f"Абзац {index} главы {number}.", ("top", "bottom") + SIDES)
            for index in range(1, size + 1)]


def classic_box(number: int, size: int = 3):
    """Глава, у которой верх только у первого абзаца, а низ у последнего."""
    rows = []
    for index in range(1, size + 1):
        sides = SIDES
        if index == 1:
            sides = ("top",) + sides
        if index == size:
            sides = sides + ("bottom",)
        rows.append((f"Абзац {index} главы {number}.", sides))
    return rows


class BlocksTestCase(unittest.TestCase):
    def setUp(self):
        self.tmpdir = TemporaryDirectory()
        self.tmp = Path(self.tmpdir.name)
        self.addCleanup(self.tmpdir.cleanup)

    def doc(self, rows, name="book.docx") -> Path:
        return make_docx(self.tmp / name, rows)

    def texts(self, path, way="auto"):
        found, used = blocks.chapters(path, way)
        return [chapter.paragraphs for chapter in found], used


class TestBoxesAroundEveryParagraph(BlocksTestCase):
    """Настоящий случай: рамка стоит на каждом абзаце по отдельности.

    Word склеивает соседние абзацы с одинаковой рамкой в один прямоугольник
    и не рисует между ними черты. Значит, двести таких абзацев подряд — это
    одна глава, а не двести.
    """

    def test_paragraphs_in_one_frame_stay_one_chapter(self):
        """Пять абзацев с одинаковой рамкой — одна глава, а не пять."""
        rows = chapter_box(1, size=5) + [("", ())] + chapter_box(2, size=2)
        found, _ = self.texts(self.doc(rows), "boxes")
        self.assertEqual([len(group) for group in found], [5, 2])

    def test_a_blank_paragraph_between_frames_starts_a_chapter(self):
        rows = chapter_box(1) + [("", ())] + chapter_box(2) + [("", ())] \
            + chapter_box(3)
        found, way = self.texts(self.doc(rows), "boxes")
        self.assertEqual(way, "boxes")
        self.assertEqual(len(found), 3)
        for number, paragraphs in enumerate(found, 1):
            self.assertTrue(all(f"главы {number}." in text for text in paragraphs),
                            paragraphs)

    def test_a_blank_paragraph_is_not_a_chapter_of_its_own(self):
        rows = chapter_box(1) + [("", ()), ("", ())] + chapter_box(2)
        found, _ = self.texts(self.doc(rows), "boxes")
        self.assertEqual(len(found), 2)


class TestClassicFrame(BlocksTestCase):
    """Рамка, у которой верх у первого абзаца, а низ у последнего."""

    def test_the_frame_holds_together(self):
        rows = classic_box(1) + [("", ())] + classic_box(2)
        found, _ = self.texts(self.doc(rows), "boxes")
        self.assertEqual([len(group) for group in found], [3, 3])

    def test_frames_touching_without_a_blank_still_split(self):
        """Низ одной рамки вплотную к верху другой — Word чертит линию."""
        found, _ = self.texts(self.doc(classic_box(1) + classic_box(2)), "boxes")
        self.assertEqual([len(group) for group in found], [3, 3])


class TestTextOutsideAnyFrame(BlocksTestCase):
    def test_it_joins_the_chapter_before_it(self):
        """«Читать далее» отдельным файлом на диск лечь не должно."""
        rows = chapter_box(1) + [("Читать далее", ())] + chapter_box(2)
        found, _ = self.texts(self.doc(rows), "boxes")
        self.assertEqual(len(found), 2)
        self.assertIn("Читать далее", found[0])

    def test_a_document_without_any_frames_says_so(self):
        rows = [(f"Просто абзац {index}.", ()) for index in range(1, 6)]
        with self.assertRaises(NoBlocksFound):
            blocks.chapters(self.doc(rows), "boxes")


class TestBlankParagraphs(BlocksTestCase):
    def test_every_blank_starts_a_chapter(self):
        rows = [("Первая.", ()), ("", ()), ("Вторая.", ()), ("", ()),
                ("Третья.", ())]
        found, way = self.texts(self.doc(rows), "blank")
        self.assertEqual(way, "blank")
        self.assertEqual(found, [["Первая."], ["Вторая."], ["Третья."]])

    def test_several_blanks_in_a_row_mean_the_same_as_one(self):
        rows = [("Первая.", ()), ("", ()), ("", ()), ("", ()), ("Вторая.", ())]
        found, _ = self.texts(self.doc(rows), "blank")
        self.assertEqual(len(found), 2)

    def test_a_document_without_blanks_says_so(self):
        rows = [("Первая.", ()), ("Вторая.", ())]
        with self.assertRaises(NoBlocksFound):
            blocks.chapters(self.doc(rows), "blank")


class TestTables(BlocksTestCase):
    """Тот же блок с рамкой Word иногда переносит таблицей в одну ячейку."""

    def table_doc(self, count: int) -> Path:
        from docx import Document

        document = Document()
        for number in range(1, count + 1):
            table = document.add_table(rows=1, cols=1)
            cell = table.rows[0].cells[0]
            cell.paragraphs[0].text = f"Первый абзац главы {number}."
            cell.add_paragraph(f"Второй абзац главы {number}.")
            document.add_paragraph("")
        path = self.tmp / "tables.docx"
        document.save(str(path))
        return path

    def test_each_table_is_one_chapter(self):
        found, _ = self.texts(self.table_doc(3), "boxes")
        self.assertEqual([len(group) for group in found], [2, 2, 2])

    def test_the_reader_would_have_lost_them_entirely(self):
        """Содержимое таблиц не попадает в `document.paragraphs` вовсе."""
        from docx import Document

        document = Document(str(self.table_doc(3)))
        self.assertEqual([p.text for p in document.paragraphs if p.text.strip()], [])


class TestChoosingTheWayItself(BlocksTestCase):
    def test_auto_prefers_frames(self):
        rows = chapter_box(1) + [("", ())] + chapter_box(2)
        _, way = self.texts(self.doc(rows), "auto")
        self.assertEqual(way, "boxes")

    def test_auto_falls_back_to_blank_paragraphs(self):
        rows = [("Первая.", ()), ("", ()), ("Вторая.", ())]
        _, way = self.texts(self.doc(rows), "auto")
        self.assertEqual(way, "blank")

    def test_auto_says_so_when_nothing_divides_the_book(self):
        rows = [("Первая.", ()), ("Вторая.", ())]
        with self.assertRaises(NoBlocksFound):
            blocks.chapters(self.doc(rows), "auto")

    def test_one_chapter_is_not_a_division(self):
        """Книга целиком одной рамкой — это не деление, а исходный файл."""
        with self.assertRaises(NoBlocksFound):
            blocks.chapters(self.doc(chapter_box(1, size=9)), "auto")


class TestTitleFromTheMarkdownHeading(BlocksTestCase):
    """Со страниц, отдающих текст в Markdown, приезжает строка «# Название»."""

    def head(self, title: str):
        return [(f"# {title}", ("top", "bottom") + SIDES)]

    def test_the_heading_becomes_the_title(self):
        rows = self.head("Возвращение на перекрёсток") + chapter_box(1) \
            + [("", ())] + chapter_box(2)
        found, _ = blocks.chapters(self.doc(rows), "boxes")
        self.assertEqual(found[0].title, "Возвращение на перекрёсток")

    def test_the_heading_leaves_the_text(self):
        """Иначе решётка попала бы и в имя файла, и первой строкой в текст."""
        rows = self.head("Возвращение") + chapter_box(1) + [("", ())] \
            + chapter_box(2)
        found, _ = blocks.chapters(self.doc(rows), "boxes")
        self.assertFalse(any(text.startswith("#") for text in found[0].paragraphs))

    def test_chapters_without_a_heading_keep_all_their_text(self):
        """Решётка есть не у всех глав — у остальных ничего не отрезаем."""
        rows = chapter_box(1) + [("", ())] \
            + [("# Вторая", ("top", "bottom") + SIDES)] + chapter_box(2)
        found, _ = blocks.chapters(self.doc(rows), "boxes")
        self.assertEqual(found[0].title, "")
        self.assertEqual(len(found[0].paragraphs), 3)
        self.assertEqual(found[1].title, "Вторая")


class TestReadingTheDocument(BlocksTestCase):
    def test_blank_paragraphs_survive_the_reading(self):
        """Обычный читатель их выбрасывает — а здесь они и есть граница."""
        rows = [("Первая.", ()), ("", ()), ("Вторая.", ())]
        found = blocks.lines(self.doc(rows))
        self.assertEqual([row.text for row in found], ["Первая.", "", "Вторая."])

    def test_a_broken_file_is_named_in_the_error(self):
        from core.readers.base import ReadError

        path = self.tmp / "broken.docx"
        path.write_bytes(b"not a document at all")
        with self.assertRaises(ReadError):
            blocks.lines(path)


if __name__ == "__main__":
    unittest.main()
