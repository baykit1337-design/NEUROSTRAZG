"""Прямая речь в кавычках — прямой речью через тире.

Главное здесь — не переписать побольше, а не тронуть лишнего: кавычка
стоит не только вокруг реплики, и слети она с названия книги, чинить это
пришлось бы вручную по всей книге.
"""

from __future__ import annotations

import sys
import unittest
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))

from ops import speech  # noqa: E402


class TestOneLine(unittest.TestCase):
    """Ровно те две строки, с которых всё началось, и их родня."""

    def test_a_whole_line_in_quotes_becomes_a_line_with_a_dash(self):
        self.assertEqual(speech.dashed("«Я-я в порядке...♥»"),
                         "— Я-я в порядке...♥")

    def test_the_dot_outside_the_quote_stays_at_the_end(self):
        """«Быстрее». — точка стояла за кавычкой, и после снятия кавычки
        она оказывается там, где ей и место."""
        self.assertEqual(speech.dashed("«Быстрее»."), "— Быстрее.")

    def test_punctuation_inside_the_quote_survives(self):
        self.assertEqual(speech.dashed("«Быстрее!»"), "— Быстрее!")

    def test_the_authors_words_after_the_reply_stay(self):
        self.assertEqual(speech.dashed("«Что?» — спросил он."),
                         "— Что? — спросил он.")

    def test_the_dash_is_an_em_dash(self):
        """Дефис на этом месте загрузчик покажет дефисом посреди строки."""
        self.assertTrue(speech.dashed("«Да».").startswith("— "))


class TestWhatIsNotAReply(unittest.TestCase):
    """Кавычка стоит не только вокруг реплики."""

    def test_a_title_inside_a_line_is_left_alone(self):
        line = "Он читал «Войну и мир»."
        self.assertEqual(speech.dashed(line), line)

    def test_a_reply_after_the_authors_words_is_left_alone(self):
        """«Он сказал: «Быстрее».» — реплика тут не абзац, а часть
        предложения: тире вместо кавычек сломало бы фразу."""
        line = "Он сказал: «Быстрее»."
        self.assertEqual(speech.dashed(line), line)

    def test_an_unclosed_quote_is_left_alone(self):
        line = "«не закрыта"
        self.assertEqual(speech.dashed(line), line)

    def test_empty_quotes_are_left_alone(self):
        self.assertEqual(speech.dashed("«»"), "«»")

    def test_a_line_already_written_with_a_dash_does_not_change(self):
        line = "— уже через тире"
        self.assertEqual(speech.dashed(line), line)

    def test_a_plain_line_does_not_change(self):
        line = "Обычная строка."
        self.assertEqual(speech.dashed(line), line)

    def test_an_empty_line_does_not_break_anything(self):
        self.assertEqual(speech.dashed(""), "")

    def test_only_the_first_closing_quote_is_taken(self):
        """Снимаем свою кавычку, а не все подряд: то, что в кавычках
        дальше по строке, кавычками и остаётся."""
        self.assertEqual(speech.dashed("«Да», — он кивнул на «Титаник»."),
                         "— Да, — он кивнул на «Титаник».")


class TestOtherQuoteMarks(unittest.TestCase):
    """Ёлочки — главное, но лапки и прямая кавычка попадаются в сливах."""

    def test_curly_quotes_work_the_same(self):
        self.assertEqual(speech.dashed("“Быстрее”."), "— Быстрее.")

    def test_straight_quotes_work_the_same(self):
        """У прямой кавычки открывающая и закрывающая — один знак."""
        self.assertEqual(speech.dashed('"Быстрее".'), "— Быстрее.")


def book():
    """Две главы: в первой речь в кавычках, во второй — уже нет."""
    return [
        ("Глава 1", ["«Я-я в порядке...♥»", "«Быстрее».",
                     "Он читал «Войну и мир»."]),
        ("Глава 2", ["— Уже через тире.", "Обычная строка."]),
    ]


class TestLookingAtTheBook(unittest.TestCase):

    def test_it_counts_only_what_will_change(self):
        found = speech.inspect(book())
        self.assertEqual(found.changed, 2)

    def test_it_counts_every_chapter(self):
        self.assertEqual(speech.inspect(book()).chapters, 2)

    def test_every_change_shows_both_sides(self):
        found = speech.inspect(book())
        self.assertEqual(found.samples[0].before, "«Я-я в порядке...♥»")
        self.assertEqual(found.samples[0].after, "— Я-я в порядке...♥")

    def test_the_change_says_which_chapter_it_is_in(self):
        self.assertEqual(speech.inspect(book()).samples[0].chapter, "Глава 1")

    def test_a_book_without_quoted_speech_is_clean(self):
        found = speech.inspect([("Глава 1", ["— Уже через тире."])])
        self.assertTrue(found.clean)
        self.assertIn("не нашлось", found.summary())

    def test_the_summary_says_how_many(self):
        self.assertIn("2", speech.inspect(book()).summary())

    def test_a_long_book_says_how_many_are_left_unshown(self):
        many = [("Глава 1", [f"«Реплика {n}»." for n in range(speech.SHOW + 5)])]
        said = speech.inspect(many).as_dict()
        self.assertEqual(len(said["samples"]), speech.SHOW)
        self.assertEqual(said["more"], 5)


class TestRewritingTheBook(unittest.TestCase):

    def test_the_replies_come_out_with_a_dash(self):
        made, _ = speech.rewrite(book())
        self.assertEqual(made[0][1][:2], ["— Я-я в порядке...♥", "— Быстрее."])

    def test_what_was_not_a_reply_stays_word_for_word(self):
        made, _ = speech.rewrite(book())
        self.assertEqual(made[0][1][2], "Он читал «Войну и мир».")
        self.assertEqual(made[1][1], ["— Уже через тире.", "Обычная строка."])

    def test_it_says_how_many_it_changed(self):
        _, count = speech.rewrite(book())
        self.assertEqual(count, 2)

    def test_the_chapters_and_their_titles_survive(self):
        made, _ = speech.rewrite(book())
        self.assertEqual([title for title, _ in made], ["Глава 1", "Глава 2"])

    def test_the_source_chapters_are_not_touched(self):
        """Переписанное пишется рядом, а не поверх."""
        was = book()
        speech.rewrite(was)
        self.assertEqual(was[0][1][0], "«Я-я в порядке...♥»")

    def test_looking_and_rewriting_agree(self):
        """Разойдись они, в списке значилось бы одно, а в книгу легло бы
        другое."""
        found = speech.inspect(book())
        _, count = speech.rewrite(book())
        self.assertEqual(found.changed, count)




class FileBase(unittest.TestCase):
    """Работа берёт файлы любого формата, а не только собранную книгу."""

    BODY = ["«Я-я в порядке...♥»", "«Быстрее».", "«Что?» — спросил он.",
            "Он читал «Войну и мир»."]

    def setUp(self):
        from tempfile import TemporaryDirectory

        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)
        self.src = self.tmp / "исходники"
        self.src.mkdir()

    def put(self, suffix=".docx", name="Глава 1", body=None):
        from core import formats
        from core.models import Chapter

        path = self.src / f"{name}{suffix}"
        formats.write(path, [Chapter(number=1, title="Глава 1",
                                     paragraphs=list(body or self.BODY))],
                      headings=True)
        return path

    def book(self, name="книга.md"):
        """Готовая книга для загрузчика — с ценой и томом в заголовке."""
        path = self.src / name
        path.write_text(" # [Глава 7 :|: 3 :|: 1 :|: Первый ]\n"
                        "«Быстрее».\n"
                        "Он читал «Войну и мир».\n", encoding="utf-8")
        return path

    def lines(self, path):
        from core import formats

        return [line for chapter in formats.read(path)
                for line in chapter.paragraphs]


class TestLookingAtFiles(FileBase):

    def test_it_reads_a_word_file(self):
        """С этого всё и началось: файл у человека — вордовский."""
        self.put(".docx")
        found = speech.look([str(self.src)])
        self.assertEqual(found.changed, 3)

    def test_every_readable_format_goes_through(self):
        for suffix in (".txt", ".docx", ".rtf", ".odt", ".fb2"):
            with self.subTest(suffix):
                self.setUp()
                self.put(suffix)
                self.assertEqual(speech.look([str(self.src)]).changed, 3)

    def test_a_whole_folder_at_once(self):
        for n in (1, 2, 3):
            self.put(".txt", name=f"Глава {n}")
        found = speech.look([str(self.src)])
        self.assertEqual(found.files, 3)
        self.assertEqual(found.changed, 9)

    def test_a_change_says_which_file_it_is_in(self):
        self.put(".txt", name="ОРИГ ЛАБИРИНТ")
        found = speech.look([str(self.src)])
        self.assertEqual(found.samples[0].file, "ОРИГ ЛАБИРИНТ.txt")

    def test_an_unreadable_file_does_not_stop_the_rest(self):
        """Папку берут целиком, и один сломанный файл в ней — не повод
        бросать остальные."""
        self.put(".txt")
        (self.src / "сломан.docx").write_bytes(b"not a zip, not a docx")
        found = speech.look([str(self.src)])
        self.assertEqual(found.changed, 3)
        self.assertTrue(found.unreadable)

    def test_nothing_is_written_while_looking(self):
        self.put(".txt")
        before = sorted(p.name for p in self.src.iterdir())
        speech.look([str(self.src)])
        self.assertEqual(sorted(p.name for p in self.src.iterdir()), before)


class TestRewritingFiles(FileBase):

    def out(self):
        return self.tmp / "готово"

    def test_the_word_file_comes_out_a_word_file(self):
        """Работа правит текст, а не перегоняет книгу из формата в
        формат: принесли .docx — получите .docx."""
        self.put(".docx")
        speech.run([str(self.src)], self.out())
        self.assertEqual([p.name for p in self.out().iterdir()],
                         ["Глава 1.docx"])

    def test_the_replies_come_out_with_a_dash(self):
        self.put(".docx")
        speech.run([str(self.src)], self.out())
        self.assertEqual(self.lines(self.out() / "Глава 1.docx")[:2],
                         ["— Я-я в порядке...♥", "— Быстрее."])

    def test_what_was_not_a_reply_survives(self):
        self.put(".docx")
        speech.run([str(self.src)], self.out())
        self.assertIn("Он читал «Войну и мир».",
                      self.lines(self.out() / "Глава 1.docx"))

    def test_the_originals_are_left_alone(self):
        path = self.put(".txt")
        was = path.read_bytes()
        speech.run([str(self.src)], self.out())
        self.assertEqual(path.read_bytes(), was)

    def test_it_says_how_many_replies_it_rewrote(self):
        self.put(".txt")
        report = speech.run([str(self.src)], self.out())
        self.assertEqual(report.extra["changed"], 3)
        self.assertEqual(report.written, 1)

    def test_the_format_can_be_asked_for(self):
        self.put(".docx")
        speech.run([str(self.src)], self.out(), out_format=".txt")
        self.assertEqual([p.name for p in self.out().iterdir()],
                         ["Глава 1.txt"])

    def test_two_files_of_one_name_do_not_overwrite_each_other(self):
        """В папке лежат «Глава 1.txt» и «Глава 1.docx». Приведи мы оба к
        одному формату — второй затёр бы первый."""
        self.put(".txt")
        self.put(".docx")
        speech.run([str(self.src)], self.out(), out_format=".txt")
        self.assertEqual(len(list(self.out().iterdir())), 2)

    def test_a_broken_file_is_reported_and_the_rest_are_written(self):
        self.put(".txt")
        (self.src / "сломан.docx").write_bytes(b"not a zip, not a docx")
        report = speech.run([str(self.src)], self.out())
        self.assertEqual(report.written, 1)
        self.assertTrue(report.failures)


class TestAReadyBookKeepsItsHeaders(FileBase):
    """Книга для загрузчика — особый случай.

    Прочитай её обычным читателем, и строки `# [Название :|: …]` стали бы
    обычным текстом, а при записи первая из них поехала бы. Цена и том
    живут в той же строке: пересобрать её значило бы поменять книге цену.
    """

    def test_the_header_survives_word_for_word(self):
        self.book()
        speech.run([str(self.src)], self.tmp / "готово")
        text = (self.tmp / "готово" / "книга.md").read_text(encoding="utf-8")
        self.assertIn(" # [Глава 7 :|: 3 :|: 1 :|: Первый ]", text)

    def test_the_reply_is_still_rewritten(self):
        self.book()
        speech.run([str(self.src)], self.tmp / "готово")
        text = (self.tmp / "готово" / "книга.md").read_text(encoding="utf-8")
        self.assertIn("— Быстрее.", text)
        self.assertIn("Он читал «Войну и мир».", text)

    def test_the_header_is_not_counted_as_a_reply(self):
        self.book()
        self.assertEqual(speech.look([str(self.src)]).changed, 1)

    def test_a_plain_markdown_is_not_mistaken_for_a_book(self):
        """У обычного `.md` заголовков загрузчика нет, и разбирать его
        как книгу нельзя: половина текста ушла бы в никуда."""
        (self.src / "просто.md").write_text("«Быстрее».\nОбычная строка.\n",
                                            encoding="utf-8")
        self.assertEqual(speech.look([str(self.src)]).changed, 1)


class TestNothingAppearsThatWasNotThere(FileBase):
    """Работа правит речь и больше ничего.

    Здесь была беда: у вордовского документа без стилей заголовка
    название берётся из имени файла, а запись «с заголовками» вставляла
    его в текст новой строкой. Человек открывал готовый файл и видел
    наверху «ОРИГ ЛАБИРИНТ 80-200» — строку, которой он не писал.
    """

    def word(self, name, heading="", body=None):
        from docx import Document

        doc = Document()
        if heading:
            doc.add_heading(heading, level=1)
        for line in (body or ["«Быстрее».", "Обычная строка."]):
            doc.add_paragraph(line)
        path = self.src / f"{name}.docx"
        doc.save(str(path))
        return path

    def written(self, name):
        """Абзацы готового файла — как их видит Word, а не читалка.

        Читалка добавленный заголовок как раз прячет: снимает его
        обратно в название. Смотреть надо в сам документ, иначе беда
        остаётся невидимой."""
        from docx import Document

        return [(p.style.name, p.text)
                for p in Document(str(self.tmp / "готово" / f"{name}.docx"))
                .paragraphs]

    def test_a_file_without_a_heading_does_not_get_one(self):
        self.word("ОРИГ ЛАБИРИНТ 80-200")
        speech.run([str(self.src)], self.tmp / "готово")
        made = self.written("ОРИГ ЛАБИРИНТ 80-200")
        self.assertTrue(all("Heading" not in style for style, _ in made), made)
        self.assertNotIn("ОРИГ ЛАБИРИНТ 80-200", [text for _, text in made])

    def test_the_text_keeps_exactly_its_own_lines(self):
        self.word("ОРИГ ЛАБИРИНТ 80-200")
        speech.run([str(self.src)], self.tmp / "готово")
        self.assertEqual([text for _, text in
                          self.written("ОРИГ ЛАБИРИНТ 80-200")],
                         ["— Быстрее.", "Обычная строка."])

    def test_a_real_heading_survives(self):
        """Заголовок, стоявший в файле, обязан вернуться на место —
        иначе «не добавляем лишнего» превратилось бы в «теряем своё»."""
        self.word("Глава 7", heading="Глава 7 — Пурпурная молния")
        speech.run([str(self.src)], self.tmp / "готово")
        made = self.written("Глава 7")
        self.assertEqual(made[0][1], "Глава 7 — Пурпурная молния")
        self.assertIn("Heading", made[0][0])

    def test_a_plain_file_whose_first_line_is_the_heading_keeps_it(self):
        """У `.txt` заголовок — обычная первая строка, и читалка забирает
        её в название. Вернуть её обязаны: она была в тексте."""
        path = self.src / "Глава 5.txt"
        path.write_text("Глава 5\n\n«Быстрее».\n", encoding="utf-8")
        speech.run([str(self.src)], self.tmp / "готово")
        said = (self.tmp / "готово" / "Глава 5.txt").read_text(encoding="utf-8")
        self.assertEqual(said.count("Глава 5"), 1, said)
        self.assertIn("— Быстрее.", said)

    def test_a_plain_file_without_a_heading_does_not_get_one(self):
        path = self.src / "ОРИГ ЛАБИРИНТ.txt"
        path.write_text("«Быстрее».\n\nОбычная строка.\n", encoding="utf-8")
        speech.run([str(self.src)], self.tmp / "готово")
        said = (self.tmp / "готово" / "ОРИГ ЛАБИРИНТ.txt").read_text(
            encoding="utf-8")
        self.assertNotIn("ОРИГ ЛАБИРИНТ", said, said)


class TestWhereTheTitleCameFrom(unittest.TestCase):
    """Читалка говорит, стоял ли заголовок в самом тексте.

    Без этого записать файл обратно нельзя: имя файла и настоящий
    заголовок часто совпадают, и по одному их названию не различить.
    """

    def setUp(self):
        from tempfile import TemporaryDirectory

        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)

    def read(self, name, text):
        from core import formats

        path = self.tmp / name
        path.write_text(text, encoding="utf-8")
        return formats.read(path)[0]

    def test_a_first_line_heading_is_from_the_text(self):
        chapter = self.read("Глава 5.txt", "Глава 5\n\nТекст.\n")
        self.assertTrue(chapter.heading_from_text)

    def test_a_name_taken_from_the_file_is_not(self):
        chapter = self.read("ОРИГ ЛАБИРИНТ.txt", "Просто текст.\n")
        self.assertFalse(chapter.heading_from_text)

    def test_the_word_document_says_it_too(self):
        from docx import Document

        from core import formats

        for heading, expected in (("Глава 7", True), ("", False)):
            with self.subTest(heading=heading or "без заголовка"):
                doc = Document()
                if heading:
                    doc.add_heading(heading, level=1)
                doc.add_paragraph("Текст.")
                path = self.tmp / "Файл.docx"
                doc.save(str(path))
                self.assertEqual(formats.read(path)[0].heading_from_text,
                                 expected)


class TestOverHttp(unittest.TestCase):
    """Карточка «Речь в кавычках» на вкладке «Инструменты»."""

    def setUp(self):
        from tempfile import TemporaryDirectory

        from core import formats
        from core.models import Chapter
        from webapp import app as web

        self._dir = TemporaryDirectory()
        self.addCleanup(self._dir.cleanup)
        self.tmp = Path(self._dir.name)
        self.src = self.tmp / "исходники"
        self.src.mkdir()
        formats.write(self.src / "Глава 1.docx",
                      [Chapter(number=1, title="Глава 1",
                               paragraphs=["«Быстрее».", "Обычная строка."])],
                      headings=True)
        web.app.config["TESTING"] = True
        self.app = web.app.test_client()
        self.web = web

    def test_the_preview_says_what_will_change(self):
        said = self.app.post("/api/speech/preview",
                             json={"targets": [str(self.src)]}).get_json()
        self.assertEqual(said["changed"], 1)
        self.assertEqual(said["samples"][0]["after"], "— Быстрее.")

    def test_the_preview_needs_files(self):
        res = self.app.post("/api/speech/preview", json={"targets": []})
        self.assertEqual(res.status_code, 400)

    def test_the_run_writes_a_new_folder(self):
        res = self.app.post("/api/speech/start", json={
            "targets": [str(self.src)], "base": str(self.tmp),
            "folder": "Речь"})
        self.assertEqual(res.status_code, 200, res.get_json())
        job_id = res.get_json()["job"]["id"]
        self.web.JOBS[job_id].thread.join(timeout=60)
        job = self.web.JOBS[job_id]
        self.assertIsNone(job.error)
        self.assertEqual(job.report["changed"], 1)
        self.assertTrue((self.tmp / "Речь" / "Глава 1.docx").is_file())

    def test_a_run_without_a_folder_is_refused(self):
        res = self.app.post("/api/speech/start",
                            json={"targets": [str(self.src)],
                                  "base": str(self.tmp)})
        self.assertEqual(res.status_code, 400)

    def test_the_old_place_is_gone(self):
        """Работа переехала в «Инструменты»: два входа в неё означали бы
        два разных ответа на один вопрос."""
        res = self.app.post("/api/format/speech",
                            json={"targets": [str(self.src)]})
        self.assertEqual(res.status_code, 404)


if __name__ == "__main__":
    unittest.main()
